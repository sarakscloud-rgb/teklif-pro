import streamlit as st
import pandas as pd
import sqlite3
import base64
import io
import sys
import subprocess
import requests
import os
import tempfile
import urllib3
import xml.etree.ElementTree as ET
import time
import shutil
from datetime import datetime
from datetime import date
from PIL import Image
from xlsxwriter.utility import xl_rowcol_to_cell

def sayiyi_yaziya_cevir(sayi):
    if sayi == 0: return "SIFIR"
    
    birler = ["", "BİR", "İKİ", "ÜÇ", "DÖRT", "BEŞ", "ALTI", "YEDİ", "SEKİZ", "DOKUZ"]
    onlar = ["", "ON", "YİRMİ", "OTUZ", "KIRK", "ELLİ", "ALTMIŞ", "YETMİŞ", "SEKSEN", "DOKSAN"]
    basamaklar = ["", "BİN", "MİLYON", "MİLYAR"]
    
    sayi_str = str(int(sayi))
    if len(sayi_str) > 12: return "ÇOK BÜYÜK SAYI"
    
    # Gruplara ayır (sağdan 3'erli)
    gruplar = []
    while sayi_str:
        gruplar.append(sayi_str[-3:])
        sayi_str = sayi_str[:-3]
    
    yazi = []
    for i, grup in enumerate(gruplar):
        grup_sayi = int(grup)
        if grup_sayi == 0: continue
        
        grup_yazi = []
        yuzler = grup_sayi // 100
        onluk = (grup_sayi % 100) // 10
        birlik = grup_sayi % 10
        
        if yuzler:
            if yuzler > 1: grup_yazi.append(birler[yuzler])
            grup_yazi.append("YÜZ")
        if onluk:
            grup_yazi.append(onlar[onluk])
        if birlik:
            if i == 1 and grup_sayi == 1: pass # "Bir Bin" denmez, sadece "Bin" denir
            else: grup_yazi.append(birler[birlik])
            
        # Basamak ismi (Bin, Milyon)
        if i > 0:
            grup_yazi.append(basamaklar[i])
            
        yazi.append(" ".join(grup_yazi))
        
    return " ".join(reversed(yazi))

# SSL Uyarılarını Gizle
urllib3.disable_warnings(urllib3.exceptions.InsecureRequestWarning)

# ==============================================================================
# 1. SAYFA VE GENEL AYARLAR
# ==============================================================================
st.set_page_config(
    page_title="Saraks Mobilya - Profesyonel Teklif Sistemi", 
    layout="wide",
    page_icon="🪚",
    initial_sidebar_state="expanded"
)

# ==============================================================================
# 2. KÜTÜPHANE KONTROL VE YÜKLEME
# ==============================================================================
# Gerekli kütüphanelerin listesi
required_packages = [
    "streamlit-cropper", 
    "requests", 
    "fpdf2", 
    "streamlit-aggrid", 
    "xlsxwriter"
]

# Kütüphaneleri kontrol et ve yoksa yükle
for package in required_packages:
    try:
        if package == "fpdf2":
            __import__("fpdf")
        elif package == "streamlit-aggrid":
            __import__("st_aggrid")
        elif package == "xlsxwriter":
            __import__("xlsxwriter")
        else:
            __import__(package.replace("-", "_"))
    except ImportError:
        st.warning(f"⚠️ Gerekli kütüphane ({package}) eksik. Otomatik yükleniyor, lütfen bekleyin...")
        try:
            subprocess.check_call([sys.executable, "-m", "pip", "install", package])
            st.success(f"✅ {package} başarıyla kuruldu! Uygulama yeniden başlatılıyor...")
            time.sleep(1)
            st.rerun()
        except Exception as e:
            st.error(f"Kütüphane yüklenirken hata oluştu: {e}")
            st.stop()

# Başarılı yükleme sonrası importlar
from streamlit_cropper import st_cropper
from fpdf import FPDF
import xlsxwriter
from st_aggrid import AgGrid
from st_aggrid import GridOptionsBuilder
from st_aggrid import GridUpdateMode
from st_aggrid import DataReturnMode
from st_aggrid import JsCode

# ==============================================================================
# SÖZLEŞME SİSTEMİ VERİTABANI VE KAYIT FONKSİYONLARI (ÜST KISMA TAŞINACAK)
# ==============================================================================

def sozlesme_tablosu_olustur():
    conn = sqlite3.connect('teklif_yonetim_sistemi.db')
    c = conn.cursor()
    c.execute('''CREATE TABLE IF NOT EXISTS sozlesmeler
                 (id INTEGER PRIMARY KEY AUTOINCREMENT,
                  firma_adi TEXT,
                  proje_adi TEXT,
                  tarih TEXT,
                  dosya_yolu TEXT,
                  tutar TEXT)''')
    conn.commit()
    conn.close()

# Tabloyu her açılışta kontrol etmesi için çağırıyoruz
sozlesme_tablosu_olustur()

def sozlesme_kaydet(firma, proje, tarih, yol, tutar):
    conn = sqlite3.connect('teklif_yonetim_sistemi.db')
    c = conn.cursor()
    c.execute("INSERT INTO sozlesmeler (firma_adi, proje_adi, tarih, dosya_yolu, tutar) VALUES (?, ?, ?, ?, ?)",
              (firma, proje, tarih, yol, tutar))
    conn.commit()
    conn.close()

def sozlesme_sil(sozlesme_id):
    conn = sqlite3.connect('teklif_yonetim_sistemi.db')
    c = conn.cursor()
    c.execute("DELETE FROM sozlesmeler WHERE id = ?", (sozlesme_id,))
    conn.commit()
    conn.close()

def sozlesme_guncelle(sozlesme_id, yeni_firma, yeni_proje, yeni_tutar):
    conn = sqlite3.connect('teklif_yonetim_sistemi.db')
    c = conn.cursor()
    c.execute("UPDATE sozlesmeler SET firma_adi = ?, proje_adi = ?, tutar = ? WHERE id = ?", 
              (yeni_firma, yeni_proje, yeni_tutar, sozlesme_id))
    conn.commit()
    conn.close()

def sozlesmeleri_getir():
    conn = sqlite3.connect('teklif_yonetim_sistemi.db')
    try:
        # Pandas ile sözleşme listesini çekiyoruz
        df = pd.read_sql("SELECT * FROM sozlesmeler ORDER BY id DESC", conn)
    except:
        df = pd.DataFrame()
    conn.close()
    return df

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_contract_docx(
    proje_adi, toplam_tutar, para_birimi,
    mus_adi, mus_adres, mus_vd, 
    sozlesme_tarihi, bitis_tarihi, sehir,
    gecikme_orani, fesih_gun, garanti_suresi, odeme_plani
):
    doc = Document()
    
    # Stil Ayarları (Genel Font)
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(10)

    # 1. BAŞLIK
    heading = doc.add_paragraph('YÜKLENİCİ HİZMET SÖZLEŞMESİ')
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    heading.runs[0].bold = True
    heading.runs[0].font.size = Pt(14)
    
    # 2. GİRİŞ BİLGİLERİ (TABLO)
    # Tutar Hesaplama (Yazıyla)
    try:
        tutar_rakam = format_para(toplam_tutar, para_birimi)
        tutar_yazi = sayiyi_yaziya_cevir(toplam_tutar)
        pb_yazi = "TÜRK LİRASI" if "TL" in para_birimi or "TRY" in para_birimi else para_birimi
        tutar_komple = f"{tutar_rakam} ( {tutar_yazi} {pb_yazi} )"
    except:
        tutar_komple = f"{toplam_tutar} {para_birimi}"

    # Tablo Oluştur
    table = doc.add_table(rows=9, cols=2)
    table.style = 'Table Grid' # Çizgili tablo olsun ki düzenli dursun
    
    def tablo_satir(row_idx, baslik, deger):
        row = table.rows[row_idx]
        row.cells[0].text = baslik
        row.cells[0].paragraphs[0].runs[0].bold = True
        row.cells[1].text = str(deger)

    tablo_satir(0, "SÖZLEŞMENİN KONUSU", f"{proje_adi} işinin yapılması")
    tablo_satir(1, "YÜKLENİCİ FİRMA", "ONUR ALIŞIK – SARAKS MOBİLYA")
    tablo_satir(2, "YÜKLENİCİ ADRES", "ALTINOVA MH. 3.YAZICI SK. NO:2 OSMANGAZİ - BURSA")
    tablo_satir(3, "VERGİ DAİRESİ", "27385398522 - ULUDAĞ V.D.")
    tablo_satir(4, "İŞVEREN FİRMA", mus_adi)
    tablo_satir(5, "İŞVEREN ADRES", mus_adres)
    tablo_satir(6, "İŞVEREN VD", mus_vd)
    tablo_satir(7, "SÖZLEŞME BEDELİ", tutar_komple)
    tablo_satir(8, "TARİHLER", f"Sözleşme: {sozlesme_tarihi}  |  Bitiş: {bitis_tarihi}")

    doc.add_paragraph("\n") # Boşluk

    # --- MADDELER ---
    def madde_ekle(baslik, metin):
        p = doc.add_paragraph()
        run = p.add_run(baslik)
        run.bold = True
        p.add_run("\n" + metin)

    # Madde 1
    madde_ekle("Madde 1- Taraflar:", 
               f"Bir tarafta SARAKS MOBİLYA adına hareket eden ONUR ALIŞIK ile diğer tarafta {mus_adi} arasında aşağıdaki şartlar dahilinde bu sözleşme akdedilmiştir. "
               f"Sözleşme metninde SARAKS MOBİLYA-ONUR ALIŞIK (YÜKLENİCİ), {mus_adi} (İŞVEREN) kelimeleri ile ifade edilmiştir.")

    # Madde 2
    madde_ekle("Madde 2- Sözleşmenin Konusu:", 
               f"İŞVEREN tarafından yaptırılacak olan: {proje_adi} projesine ait işin yaptırılmasıdır. "
               f"İş Anahtar teslimi götürü fiyat esasına göre yapılacak olup işin götürü fiyat tutarı {tutar_komple}'dir. "
               "YÜKLENİCİ söz konusu işi özel şartnamedeki hususlara göre Ana Sözleşme ve Sözleşme şartlarına uygun olarak zamanında yapmayı kabul ve taahhüt etmiştir.")

    # Madde 3
    madde_ekle("Madde 3- Sözleşmenin Ekleri:", "1- Proje Teklif Dosyası\n3- Proje Uygulama Çizimleri")

    # Madde 4 (Tam Metin)
    text_m4 = (
        f"İşin müddeti “Mücbir Sebepler” dışında iş bu müddet dahilinde bitmediği takdirde geçecek beher gün için "
        f"YÜKLENİCİ’den Sözleşme Tutarının {gecikme_orani} Oranında gecikme cezası kesilecektir. "
        f"Şu kadar ki; gecikme {fesih_gun} günü geçerse İŞVEREN firma gecikme cezasını almaya devam ederek beklemeye "
        f"veya mahkeme kanalı ile tespit yaptırmaya, ayrıca bir karar almaya, protesto çekmeye lüzum kalmaksızın "
        f"YÜKLENİCİ’nin nam ve hesabına işi bir başkasına yaptırmaya veya sözleşmeyi fesih ederek, uğradığı zararları "
        f"veya üçüncü şahsa ödemeye mecbur kalacağı cezai şart ya da tazminatı YÜKLENİCİ’den talep etmeye yetkilidir."
    )
    madde_ekle("Madde 4- Müddet ve Gecikme Cezası:", text_m4)

    # Madde 5 (Tam Metin)
    text_m5 = (
        f"YÜKLENİCİ, sözleşmenin imzalanmasına müteakip 3 gün içerisinde taahhüdünün 5.maddesinde yazılı süre içerisinde "
        f"mevcut işi bitireceğini belirten bir iş programını İŞVEREN firmaya vermekle mükelleftir. "
        f"Bu iş programı işveren firmaya verilmemesi halinde gecikilen her gün için Sözleşme Tutarının {gecikme_orani} Oranında "
        f"cezayı işverene ödemekle yükümlüdür. İşveren, YÜKLENİCİ firma tarafından yapılmış iş programına işi durumunu göre "
        f"müdahil olarak ara temrinler tespit edebilir. YÜKLENİCİ bu tespitlere itiraz edemez."
    )
    madde_ekle("Madde 5- İş Programı:", text_m5)

    # Madde 6
    madde_ekle("Madde 6- Vergi ve Vesair Masraflar:", 
               "Taahhüdün ifasına ait her türlü vergi, resim ve harçlar ile Noter masrafları her çeşit sigorta primleri, işçi ve işveren hisseleri, İşsizlik Sigortası Primleri, fazla mesai, ikramiye ve Pazar yevmiyeleri gibi iş kanununun gerektirdiği bilcümle vecibeler YÜKLENİCİ’a aittir. YÜKLENİCİ mevcut vergi, resim, prim ve harçların artması veya rayiçlerin yükselmesi veya yeniden vergi, prim ve harçlar ihdası gibi sebeplere dayanarak yeni fiyat veya süre uzatılması gibi taleplerde bulunamaz.")

    # Madde 7
    madde_ekle("Madde 7- İmalatın ve İhzaratın Muhafazası:", 
               "YÜKLENİCİun kendi işyerinde veya inşaat mahallinde yaptığı ihzarat, İŞVEREN firmasının teslim ettiği malzemeler ve imalatların İŞVEREN firmasına teslimine kadar muhafazası ve mesuliyeti YÜKLENİCİ’a aittir.")

    # Madde 8
    madde_ekle("Madde 8- Kusurlu Malzeme, İmalat ve Hasarlar:", 
               "Şartnamelerde yazılı hükümlere uymayan veya fen gereği olarak belli vasıfları ve şartları haiz olmayan malzeme ile proje ve şartnamesine ve tekniğine uymayan imalat red olunur. YÜKLENİCİ kabul edilmeyen malzemeyi değiştirmeye veya düzeltmeye veya yeniden imal etmeye mecburdur. Bu yüzden hasıl olabilecek gecikmeler iş müddetinin ve ara terminlerin uzatılmasını gerektirmez. İŞVEREN firması, bu gibi kusurlu malzeme ve imalattan mütevellit maruz kalacağı ziyanı ayrıca YÜKLENİCİ’den talep etmek hakkına haizdir.")

    # Madde 9
    madde_ekle("Madde 9- İmalat Miktarının Artma veya Eksilmesi:", 
               "YÜKLENİCİ’nin Kapsamı dışında ve karşılıklı fiyat mutabakatı sağlanması kaydı ile İŞVEREN firmasının göreceği lüzum üzerine fazla veya eksik iş yaptırmaya yetkilidir. Ancak iş miktarındaki artma veya eksilme yekûn bedelin % 30 ’dan fazla veya eksik olamaz. YÜKLENİCİ işlerin bu nispet dahilinde ki artma veya eksilmeden dolayı zarar ve ziyan gibi bir talep ve itirazda bulunamaz.")

    # Madde 10
    madde_ekle("Madde 10- Taahhüdün Devri:", "YÜKLENİCİ bu sözleşme ile taahhüt ettiği işleri İŞVEREN firmasının yazılı onayını almaksızın kısmen veya tamamen başka birine devir ve temlik edemez.")

    # Madde 11
    madde_ekle("Madde 11- İmalatın Teslimi:", "YÜKLENİCİ tarafından işin ikmal edildiği İŞVEREN firmasına bildirilmesini müteakip İŞVEREN firmasınca teşkil olunacak heyet marifetiyle imalatın monte edildiği mahalde işin kontrolu yapılıp, kabul edilerek veya 9.madde uyarınca işlem yapılarak bir tutanak tanzim edilecektir. YÜKLENİCİ’nin imalat yerinden İŞVEREN firmanın iş yerine kadar olan her türlü nakliye, yükleme, boşaltma, istifleme giderleri ile bu meyanda meydana gelecek her türlü hasar ve kusurlar YÜKLENİCİ’ye aittir.")

    # Madde 12
    madde_ekle("Madde 12- Garanti Müddeti:", f"YÜKLENİCİ firma tarafından yapılan işlerin garanti süresi; aksi belirtilmediği sürece {garanti_suresi} olarak kabul edilir.")

    # Madde 13
    madde_ekle("Madde 13- Ödeme Şekli:", f"{odeme_plani}")

    # Madde 14 (Tam Metin)
    text_m14 = (
        "Herhangi bir ameliyenin yapılması dolayısı ile vukua gelebilecek kazalardan korunmak için YÜKLENİCİ iş güvenliği "
        "ve iş tüzüğü ile tespit edilen bütün tedbirleri alacak ve kazalardan korunma usul ve çarelerinin işçi ve personeline "
        "öğretecektir. YÜKLENİCİ, kazalara karşı her türlü emniyet tedbirlerini almakla mükellef olup, gerek ihmal, "
        "dikkatsizlik veya tedbirsizlikten, gerekse ehliyetsiz işçi kullanmaktan veya herhangi bir başka sebeplerle vuku "
        "bulacak kazalardan mesul olup, kazaya uğrayacak işçi, personel ve üçüncü kişilerin tedavi ve kendilerine, ailelerine "
        "verilecek tazminat, mahkeme masrafları ve sair masrafları tamamen YÜKLENİCİ’a aittir. İŞVEREN firması bu nedenle "
        "herhangi bir talep ve/veya ödemeye maruz kalırsa bu bedel YÜKLENİCİ tarafından aynen karşılanacaktır."
    )
    madde_ekle("Madde 14- Emniyet Tedbirleri:", text_m14)

    # Madde 15
    madde_ekle("Madde 15- Kanuni İkâmetgah:", f"YÜKLENİCİ ALTINOVA MH. 3.YAZICI SK. NO:2 OSMANGAZİ - BURSA adresini kanuni ikametgâh olarak göstermiş olup, bu adrese yapılacak tebligat aynı günde YÜKLENİCİ’un kendisine yapılmış sayılır.")

    # Madde 16
    madde_ekle("Madde 16- İhtilafların Halli:", "Bu sözleşmenin tatbikinden doğacak her türlü ihtilafların halli, mercii T.C. Bursa Mahkemeleri ve İcra Daireleridir.")

    # Madde 17 & 18
    madde_ekle("Madde 17- Sözleşme Ekleri & Madde 18- Tarih ve Yer:", f"Sözleşme eki olarak madde 3 de yer alan ekler sözleşmenin tamamlayıcısı olup, ayrılmaz bir parçasıdır. Bu sözleşme taraflar arasında {sozlesme_tarihi} tarihinde {sehir}'da tanzim ve imza edilmiştir.")
    
    # 3. İMZA ALANI
    doc.add_paragraph("\n\n")
    table_imza = doc.add_table(rows=3, cols=2)
    table_imza.width = Cm(16)
    
    row0 = table_imza.rows[0]
    row0.cells[0].text = "İŞVEREN FİRMA"
    row0.cells[0].paragraphs[0].runs[0].bold = True
    row0.cells[1].text = "YÜKLENİCİ FİRMA"
    row0.cells[1].paragraphs[0].runs[0].bold = True
    
    row1 = table_imza.rows[1]
    row1.cells[0].text = str(mus_adi)
    row1.cells[1].text = "SARAKS MOBİLYA - ONUR ALIŞIK"

    row2 = table_imza.rows[2]
    row2.cells[1].text = "ONUR ALIŞIK"
    row2.cells[1].paragraphs[0].runs[0].bold = True

    return doc

# ==============================================================================
# TESLİM TUTANAĞI MOTORU (TAM SÜRÜM: DB + PDF + WORD + OTO. ÇEKME)
# BU BLOK KODUN EN ÜSTÜNDE, IMPORTLARDAN HEMEN SONRA OLMALIDIR!
# ==============================================================================
import sqlite3
import pandas as pd
from fpdf import FPDF
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

# --- 1. VERİTABANI FONKSİYONLARI ---
def tutanak_tablosu_olustur():
    conn = sqlite3.connect('teklif_yonetim_sistemi.db')
    c = conn.cursor()
    c.execute('''CREATE TABLE IF NOT EXISTS teslim_tutanaklari
                 (id INTEGER PRIMARY KEY AUTOINCREMENT,
                  firma_adi TEXT,
                  proje_adi TEXT,
                  tarih TEXT,
                  dosya_yolu TEXT)''')
    conn.commit()
    conn.close()

# Tabloyu her açılışta kontrol et (İşte hata veren satır buradaydı, şimdi tanımın altında)
tutanak_tablosu_olustur()

def tutanak_kaydet(firma, proje, tarih, yol):
    conn = sqlite3.connect('teklif_yonetim_sistemi.db')
    c = conn.cursor()
    c.execute("INSERT INTO teslim_tutanaklari (firma_adi, proje_adi, tarih, dosya_yolu) VALUES (?, ?, ?, ?)",
              (firma, proje, tarih, yol))
    conn.commit()
    conn.close()

def tutanaklari_getir():
    conn = sqlite3.connect('teklif_yonetim_sistemi.db')
    try:
        df = pd.read_sql("SELECT * FROM teslim_tutanaklari ORDER BY id DESC", conn)
    except:
        df = pd.DataFrame()
    conn.close()
    return df

def tutanak_sil(tutanak_id):
    conn = sqlite3.connect('teklif_yonetim_sistemi.db')
    c = conn.cursor()
    c.execute("DELETE FROM teslim_tutanaklari WHERE id = ?", (tutanak_id,))
    conn.commit()
    conn.close()

# --- TEKLİF ÜRÜNLERİNİ ÇEKEN FONKSİYON ---
def teklif_urunlerini_getir(teklif_id):
    conn = sqlite3.connect('teklif_yonetim_sistemi.db')
    try:
        # Teklif satırlarını çekiyoruz
        query = "SELECT urun_kodu as Kod, urun_adi as Urun, miktar || ' ' || birim as Adet FROM teklif_satirlari WHERE teklif_id = ?"
        df = pd.read_sql(query, conn, params=(teklif_id,))
    except:
        df = pd.DataFrame(columns=["Kod", "Urun", "Adet"])
    conn.close()
    return df

def create_delivery_pdf(firma, proje, sozlesme_tarihi, teslim_tarihi, urun_df):
    pdf = FPDF()
    pdf.add_page()
    
    # Font Ayarı
    font_dir = os.path.join(os.environ.get('WINDIR', 'C:\\Windows'), 'Fonts')
    main_font = 'Arial'
    if os.path.exists(os.path.join(font_dir, 'arial.ttf')):
        pdf.add_font('ArialTR', '', os.path.join(font_dir, 'arial.ttf'), uni=True)
        pdf.add_font('ArialTR', 'B', os.path.join(font_dir, 'arialbd.ttf'), uni=True)
        main_font = 'ArialTR'
    else:
        pass 

    # Başlık
    pdf.set_font(main_font, 'B', 14)
    pdf.cell(0, 10, "İŞ TESLİM TUTANAĞI", 0, 1, 'C')
    pdf.ln(5)

    pdf.set_font(main_font, 'B', 10)
    
    def satir_yaz(baslik, deger):
        pdf.cell(40, 7, baslik, 0, 0)
        pdf.set_font(main_font, '', 10)
        pdf.cell(0, 7, f": {deger}", 0, 1)
        pdf.set_font(main_font, 'B', 10)

    satir_yaz("Proje Adı", proje)
    satir_yaz("İşveren Adı", firma)
    satir_yaz("Yüklenici Adı", "ONUR ALIŞIK - SARAKS MOBİLYA")
    satir_yaz("Sözleşme Tarihi", sozlesme_tarihi)
    satir_yaz("İş Teslim Tarihi", teslim_tarihi)
    
    pdf.ln(10)

    # Standart Metin
    pdf.set_font(main_font, '', 10)
    metin = (
        f"İşveren {firma} İle Yüklenici Onur Alışık arasında imzalanan sözleşme kapsamında gerçekleştirilen "
        f"ve {teslim_tarihi} tarihinde bitirilen, Yüklenicinin de hazır bulunması ile iş sahasına giderek "
        "Yüklenici tarafından yapılmış işleri Kesin Kabul bakımından incelemiş ve Ek-1’de listelenmiş ürünlerin "
        "eksiksiz, kusursuz ve montajı tamamlanmış olarak İşverene teslim edilmiştir."
    )
    pdf.multi_cell(0, 6, metin)
    
    pdf.ln(10)
    pdf.cell(0, 6, "Teslim Tutanağı 2 nüsha olarak düzenlenmiştir.", 0, 1)
    
    # İmzalar
    pdf.ln(15)
    pdf.cell(0, 6, f"Tarih: {teslim_tarihi}", 0, 1, 'R')
    pdf.ln(5)
    
    pdf.set_font(main_font, 'B', 10)
    pdf.cell(90, 6, "TESLİM EDEN (YÜKLENİCİ)", 0, 0, 'C')
    pdf.cell(90, 6, "TESLİM ALAN (İŞVEREN)", 0, 1, 'C')
    
    pdf.set_font(main_font, '', 10)
    pdf.cell(90, 6, "ONUR ALIŞIK", 0, 0, 'C')
    pdf.cell(90, 6, firma, 0, 1, 'C')
    
    # --- EK-1 SAYFASI ---
    pdf.add_page()
    pdf.set_font(main_font, 'B', 12)
    pdf.cell(0, 10, "İŞ TESLİM TUTANAĞI Ek-1", 0, 1, 'C')
    pdf.set_font(main_font, '', 10)
    pdf.cell(0, 8, "Teslim Edilen Ürünler Listesi:", 0, 1, 'L')
    pdf.ln(2)

    # Tablo Başlıkları
    pdf.set_fill_color(240, 240, 240)
    pdf.set_font(main_font, 'B', 9)
    # Sütun Genişlikleri
    pdf.cell(15, 8, "KOD", 1, 0, 'C', True)
    pdf.cell(65, 8, "ÜRÜN ADI / AÇIKLAMA", 1, 0, 'L', True)
    pdf.cell(20, 8, "ADET", 1, 0, 'C', True)
    pdf.cell(90, 8, "TESLİM NOTU", 1, 1, 'L', True)

    # Tablo İçeriği
    pdf.set_font(main_font, '', 9)
    if not urun_df.empty:
        for index, row in urun_df.iterrows():
            # --- DÜZELTME BURADA ---
            # Hem 'Ürün Adı' hem 'Urun' kontrolü yapıyoruz. Hangisi varsa onu alıyor.
            urun_adi = str(row.get('Ürün Adı', row.get('Urun', '')))
            
            pdf.cell(15, 8, str(row.get('Kod', '')), 1, 0, 'C')
            pdf.cell(65, 8, urun_adi[:55], 1, 0, 'L')  # Düzeltilmiş değişkeni kullanıyoruz
            pdf.cell(20, 8, str(row.get('Adet', '')), 1, 0, 'C')
            pdf.cell(90, 8, str(row.get('Not', '')), 1, 1, 'L')

    return pdf.output(dest='S').encode('latin-1')

def create_delivery_docx(firma, proje, sozlesme_tarihi, teslim_tarihi, urun_df):
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)

    heading = doc.add_paragraph('İŞ TESLİM TUTANAĞI')
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    heading.runs[0].bold = True
    heading.runs[0].font.size = Pt(14)

    # Bilgiler Tablosu
    table_info = doc.add_table(rows=5, cols=2)
    def set_info(idx, label, val):
        r = table_info.rows[idx]
        r.cells[0].text = label
        r.cells[0].paragraphs[0].runs[0].bold = True
        r.cells[0].width = Cm(5)
        r.cells[1].text = f": {val}"

    set_info(0, "Proje Adı", proje)
    set_info(1, "İşveren Adı", firma)
    set_info(2, "Yüklenici Adı", "ONUR ALIŞIK - SARAKS MOBİLYA")
    set_info(3, "Sözleşme Tarihi", sozlesme_tarihi)
    set_info(4, "İş Teslim Tarihi", teslim_tarihi)

    doc.add_paragraph("\n")

    p = doc.add_paragraph()
    p.add_run(
        f"İşveren {firma} İle Yüklenici Onur Alışık arasında imzalanan sözleşme kapsamında gerçekleştirilen "
        f"ve {teslim_tarihi} tarihinde bitirilen, Yüklenicinin de hazır bulunması ile iş sahasına giderek "
        "Yüklenici tarafından yapılmış işleri Kesin Kabul bakımından incelemiş ve Ek-1’de listelenmiş ürünlerin "
        "eksiksiz, kusursuz ve montajı tamamlanmış olarak İşverene teslim edilmiştir."
    ).bold = False
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    doc.add_paragraph("Teslim Tutanağı 2 nüsha olarak düzenlenmiştir.")
    doc.add_paragraph(f"Tarih: {teslim_tarihi}").alignment = WD_ALIGN_PARAGRAPH.RIGHT
    doc.add_paragraph("\n\n")

    # İmzalar
    table_sig = doc.add_table(rows=2, cols=2)
    table_sig.width = Cm(16)
    
    r0 = table_sig.rows[0]
    r0.cells[0].text = "TESLİM EDEN (YÜKLENİCİ)"
    r0.cells[0].paragraphs[0].runs[0].bold = True
    r0.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    r0.cells[1].text = "TESLİM ALAN (İŞVEREN)"
    r0.cells[1].paragraphs[0].runs[0].bold = True
    r0.cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    r1 = table_sig.rows[1]
    r1.cells[0].text = "\nONUR ALIŞIK"
    r1.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1.cells[1].text = f"\n{firma}"
    r1.cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Ek-1
    doc.add_page_break()
    h2 = doc.add_paragraph('İŞ TESLİM TUTANAĞI Ek-1')
    h2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    h2.runs[0].bold = True
    
    doc.add_paragraph("Teslim Edilen Ürünler Listesi:")

    # --- TABLO BURADA OLUŞTURULUYOR (4 SÜTUNLU) ---
    if not urun_df.empty:
        table_urun = doc.add_table(rows=1, cols=4) # Sütun sayısı 4'e çıktı
        table_urun.style = 'Table Grid'
        
        hdr = table_urun.rows[0].cells
        hdr[0].text = "KOD"; hdr[0].width = Cm(0.5)
        hdr[1].text = "ÜRÜN ADI / AÇIKLAMA"; hdr[1].width = Cm(9)
        hdr[2].text = "ADET"; hdr[2].width = Cm(2)
        hdr[3].text = "TESLİM NOTU"; hdr[3].width = Cm(4) # <--- YENİ SÜTUN
        
        for index, row in urun_df.iterrows():
            row_cells = table_urun.add_row().cells
            
            # Ürün Adı Düzeltmesi (Hem Urun hem Ürün Adı kontrolü)
            urun_adi = str(row.get('Ürün Adı', row.get('Urun', '')))
            
            row_cells[0].text = str(row.get('Kod', ''))
            row_cells[1].text = urun_adi
            row_cells[2].text = str(row.get('Adet', ''))
            row_cells[3].text = str(row.get('Not', '')) # Notu buraya yazıyoruz (Boşsa boş gelir)

    return doc

# --- 3. WORD MOTORU (4 Sütunlu: Kod, Ürün, Adet, Not) ---
def create_delivery_docx(firma, proje, sozlesme_tarihi, teslim_tarihi, urun_df):
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)

    heading = doc.add_heading('İŞ TESLİM TUTANAĞI')
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    heading.runs[0].bold = True

    # Bilgiler
    table_info = doc.add_table(rows=5, cols=2)
    def set_info(idx, label, val):
        r = table_info.rows[idx]
        r.cells[0].text = label
        r.cells[0].paragraphs[0].runs[0].bold = True
        r.cells[0].width = Cm(5)
        r.cells[1].text = f": {val}"

    set_info(0, "Proje Adı", proje)
    set_info(1, "İşveren Adı", firma)
    set_info(2, "Yüklenici Adı", "ONUR ALIŞIK - SARAKS MOBİLYA")
    set_info(3, "Sözleşme Tarihi", sozlesme_tarihi)
    set_info(4, "İş Teslim Tarihi", teslim_tarihi)

    doc.add_paragraph("\n")
    p = doc.add_paragraph(
        f"İşveren {firma} İle Yüklenici Onur Alışık arasında imzalanan sözleşme kapsamında gerçekleştirilen "
        f"ve {teslim_tarihi} tarihinde bitirilen Yüklenicinin de hazır bulunması ile iş sahasına giderek "
        "Yüklenici tarafından yapılmış işleri Kesin Kabul bakımından incelemiş ve Ek-1’de listelenmiş ürünlerin "
        "eksiksiz, kusursuz ve montajı tamamlanmış olarak İşverene teslim edilmiştir."
    )
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    doc.add_paragraph("Teslim Tutanağı 2 nüsha olarak düzenlenmiştir.")
    doc.add_paragraph(f"Tarih: {teslim_tarihi}").alignment = WD_ALIGN_PARAGRAPH.RIGHT
    doc.add_paragraph("\n\n")

    # İmzalar
    table_sig = doc.add_table(rows=2, cols=2)
    table_sig.width = Cm(16)
    r0 = table_sig.rows[0]
    r0.cells[0].text = "TESLİM EDEN (YÜKLENİCİ)"; r0.cells[0].paragraphs[0].runs[0].bold = True
    r0.cells[1].text = "TESLİM ALAN (İŞVEREN)"; r0.cells[1].paragraphs[0].runs[0].bold = True
    r1 = table_sig.rows[1]
    r1.cells[0].text = "\nONUR ALIŞIK"; r1.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1.cells[1].text = f"\n{firma}"; r1.cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # EK-1 TABLOSU
    doc.add_page_break()
    doc.add_heading('İŞ TESLİM TUTANAĞI Ek-1', level=2).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("Teslim Edilen Ürünler Listesi:")

    if not urun_df.empty:
        table_urun = doc.add_table(rows=1, cols=4)
        table_urun.style = 'Table Grid'
        
        hdr = table_urun.rows[0].cells
        hdr[0].text = "KOD"; hdr[0].width = Cm(2.5)
        hdr[1].text = "ÜRÜN ADI / AÇIKLAMA"; hdr[1].width = Cm(9)
        hdr[2].text = "ADET"; hdr[2].width = Cm(2)
        hdr[3].text = "TESLİM NOTU"; hdr[3].width = Cm(4)
        
        for index, row in urun_df.iterrows():
            row_cells = table_urun.add_row().cells
            row_cells[0].text = str(row.get('Kod', ''))
            row_cells[1].text = str(row.get('Urun', ''))
            row_cells[2].text = str(row.get('Adet', ''))
            row_cells[3].text = str(row.get('Not', ''))

    return doc

# --- 3. WORD MOTORU (TABLO DESTEKLİ) ---
def create_delivery_docx(firma, proje, sozlesme_tarihi, teslim_tarihi, urun_df):
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)

    heading = doc.add_paragraph('İŞ TESLİM TUTANAĞI')
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    heading.runs[0].bold = True
    heading.runs[0].font.size = Pt(14)

    # Bilgiler Tablosu
    table_info = doc.add_table(rows=5, cols=2)
    def set_info(idx, label, val):
        r = table_info.rows[idx]
        r.cells[0].text = label
        r.cells[0].paragraphs[0].runs[0].bold = True
        r.cells[0].width = Cm(5)
        r.cells[1].text = f": {val}"

    set_info(0, "Proje Adı", proje)
    set_info(1, "İşveren Adı", firma)
    set_info(2, "Yüklenici Adı", "ONUR ALIŞIK - SARAKS MOBİLYA")
    set_info(3, "Sözleşme Tarihi", sozlesme_tarihi)
    set_info(4, "İş Teslim Tarihi", teslim_tarihi)

    doc.add_paragraph("\n")

    p = doc.add_paragraph()
    p.add_run(
        f"İşveren {firma} İle Yüklenici Onur Alışık arasında imzalanan sözleşme kapsamında gerçekleştirilen "
        f"ve {teslim_tarihi} tarihinde bitirilen Yüklenicinin de hazır bulunması ile iş sahasına giderek "
        "Yüklenici tarafından yapılmış işleri Kesin Kabul bakımından incelemiş ve Ek-1’de listelenmiş ürünlerin "
        "eksiksiz, kusursuz ve montajı tamamlanmış olarak İşverene teslim edilmiştir."
    ).bold = False
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    doc.add_paragraph("Teslim Tutanağı 2 nüsha olarak düzenlenmiştir.")
    doc.add_paragraph(f"Tarih: {teslim_tarihi}").alignment = WD_ALIGN_PARAGRAPH.RIGHT
    doc.add_paragraph("\n\n")

    # İmzalar
    table_sig = doc.add_table(rows=2, cols=2)
    table_sig.width = Cm(16)
    
    r0 = table_sig.rows[0]
    r0.cells[0].text = "TESLİM EDEN (YÜKLENİCİ)"
    r0.cells[0].paragraphs[0].runs[0].bold = True
    r0.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    r0.cells[1].text = "TESLİM ALAN (İŞVEREN)"
    r0.cells[1].paragraphs[0].runs[0].bold = True
    r0.cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    r1 = table_sig.rows[1]
    r1.cells[0].text = "\nONUR ALIŞIK"
    r1.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1.cells[1].text = f"\n{firma}"
    r1.cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Ek-1
    doc.add_page_break()
    h2 = doc.add_paragraph('İŞ TESLİM TUTANAĞI Ek-1')
    h2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    h2.runs[0].bold = True
    
    doc.add_paragraph("Teslim Edilen Ürünler / Odalar Listesi:")

    if not urun_df.empty:
        table_urun = doc.add_table(rows=1, cols=3)
        table_urun.style = 'Table Grid'
        hdr = table_urun.rows[0].cells
        hdr[0].text = "NO"; hdr[0].width = Cm(1.5)
        hdr[1].text = "ÜRÜN ADI / AÇIKLAMA"; hdr[1].width = Cm(12)
        hdr[2].text = "ADET"; hdr[2].width = Cm(2.5)
        
        for index, row in urun_df.iterrows():
            row_cells = table_urun.add_row().cells
            row_cells[0].text = str(index + 1)
            row_cells[1].text = str(row['Ürün Adı'])
            row_cells[2].text = str(row['Adet'])

    return doc

# ==============================================================================
# 3. YARDIMCI FONKSİYONLAR
# ==============================================================================

def get_font_path(font_name, is_bold=False):
    """
    PDF oluştururken Türkçe karakter sorunu yaşamamak için 
    doğru font dosyasını (TTF) bulur.
    Önce Windows klasörüne bakar, yoksa internetten indirir.
    """
    # 1. Seçenek: Windows Font Klasörü (En Garanti Yöntem)
    windows_font_dir = os.path.join(os.environ.get('WINDIR', 'C:\\Windows'), 'Fonts')
    
    if is_bold:
        win_path = os.path.join(windows_font_dir, "arialbd.ttf")
    else:
        win_path = os.path.join(windows_font_dir, "arial.ttf")
    
    if os.path.exists(win_path):
        return win_path, "Arial"

    # 2. Seçenek: Eğer Windows fontu yoksa (Linux/Mac vb.) Roboto indir
    file_name = "Roboto-Bold.ttf" if is_bold else "Roboto-Regular.ttf"
    url = f"https://github.com/google/fonts/raw/main/apache/roboto/{file_name}"
    
    # Font dosyası yoksa veya bozuksa indir
    if not os.path.exists(file_name) or os.path.getsize(file_name) < 10000:
        try:
            r = requests.get(url, allow_redirects=True, verify=False, timeout=10)
            if r.status_code == 200:
                with open(file_name, "wb") as f:
                    f.write(r.content)
            else:
                return None, None
        except Exception as e:
            # İnternet yoksa sessizce geç
            return None, None
            
    if os.path.exists(file_name):
        return file_name, "Roboto"
    
    return None, None


def kodlari_yeniden_sirala(df, proje_kodu):
    """
    Tablodaki satır sırasına göre KOD sütununu (Örn: MUTFAK-01) yeniden oluşturur.
    Sürükle bırak işleminden sonra çağrılır.
    """
    if df is None or df.empty:
        return df
    
    # Eğer proje kodu girilmemişse varsayılan 'URUN' kullanalım
    if proje_kodu and len(str(proje_kodu).strip()) > 0:
        prefix = str(proje_kodu).strip()
    else:
        prefix = "URUN"
    
    yeni_kodlar = []
    
    for i in range(len(df)):
        # Sıra numarası (1'den başlar)
        sira = i + 1
        
        # Format: PREFİX-01 (Çift haneli sayı)
        kod = f"{prefix}-{sira:02d}"
        yeni_kodlar.append(kod)
    
    df["KOD"] = yeni_kodlar
    return df


def jpeg_icin_hazirla(img: Image.Image, arkaplan=(255, 255, 255)) -> Image.Image:
    """
    PNG formatındaki şeffaf görselleri JPEG formatına uygun hale getirir.
    Şeffaf alanları beyaz yapar.
    """
    if img.mode in ("RGBA", "LA") or (img.mode == "P" and "transparency" in img.info):
        img = img.convert("RGBA")
        bg = Image.new("RGBA", img.size, arkaplan + (255,))
        img = Image.alpha_composite(bg, img).convert("RGB")
    else:
        if img.mode != "RGB":
            img = img.convert("RGB")
    return img


def format_para(deger, sembol="TL", ondalik=2):
    """
    Sayısal değerleri para birimi formatına (1.000,00 TL) çevirir.
    """
    semboller = {
        "TL": "₺", 
        "USD": "$", 
        "EUR": "€", 
        "GBP": "£"
    }
    gercek_sembol = semboller.get(sembol, sembol)
    
    try:
        deger = float(deger)
    except:
        return f"0 {gercek_sembol}"
    
    # Formatlama mantığı: Önce İngiliz stili (1,000.00), sonra Türkçe değişimi
    format_str = f"{{:,.{ondalik}f}}"
    text = format_str.format(deger)
    
    # Nokta ve virgül değişimi
    text = text.replace(",", "X")
    text = text.replace(".", ",")
    text = text.replace("X", ".")
    
    return f"{text} {gercek_sembol}"


def temizle_ve_sayiya_cevir(df, kolonlar):
    """
    DataFrame içindeki sütunları temizleyip güvenli bir şekilde sayıya (float) çevirir.
    Özellikle '5.000,00' gibi metinleri '5000.00' sayısına dönüştürür.
    """
    def safe_convert(x):
        # Null kontrolü
        if pd.isna(x):
            return 0.0
            
        # Zaten sayıysa dokunma
        if isinstance(x, (int, float)):
            return float(x)
        
        # Stringe çevir ve boşlukları sil
        x = str(x).strip()
        if not x:
            return 0.0
        
        # --- KRİTİK KONTROL ---
        # Eğer içinde sadece nokta varsa ve virgül YOKSA (Örn: '250000.0') 
        # Bu bir Python float stringidir. Noktayı silersek sayı 10 kat büyür.
        if x.count('.') == 1 and ',' not in x:
            try:
                return float(x)
            except:
                pass 
        
        # Türkçe formatında binlik ayracı nokta, ondalık virgül ise düzelt
        # 1. Noktaları (binlik ayracı) sil
        x = x.replace(".", "")
        # 2. Virgülü (ondalık ayracı) noktaya çevir
        x = x.replace(",", ".")
        
        try:
            return float(x)
        except:
            return 0.0

    for col in kolonlar:
        if col in df.columns:
            df[col] = df[col].apply(safe_convert)
            
    return df


@st.cache_data(ttl=600) 
def kurlari_getir():
    """TCMB XML servisinden güncel döviz kurlarını çeker."""
    url = "https://www.tcmb.gov.tr/kurlar/today.xml"
    
    varsayilan_kurlar = {
        "TL": 1.0, 
        "USD": 0.0, 
        "EUR": 0.0, 
        "GBP": 0.0
    }
    
    try:
        response = requests.get(url, timeout=10, verify=False)
        if response.status_code == 200:
            root = ET.fromstring(response.content)
            for currency in root.findall('Currency'):
                code = currency.get('CurrencyCode')
                if code in ["USD", "EUR", "GBP"]:
                    rate = currency.find('ForexBuying').text 
                    if rate:
                        varsayilan_kurlar[code] = float(rate)
        return varsayilan_kurlar
    except:
        # Hata durumunda varsayılan (0) döner, arayüzde manuel giriş istenir
        return varsayilan_kurlar

# ==============================================================================
# 4. TASARIM VE CSS (DETAYLI)
# ==============================================================================
st.markdown("""
<style>
    /* --- BURASI EKLENECEK: Üst Barı ve Deploy Butonunu Gizleme --- */
    
    /* Sağ üstteki Deploy butonunu gizler */
    .stDeployButton {
        display: none;
    }
    
    /* Komple üstteki gri/beyaz şeridi (header) gizler */
    /* Böylece uygulama tam ekran gibi görünür */
    header[data-testid="stHeader"] {
        visibility: hidden;
    }
    
    /* Eğer üst bar gizlenince içerik çok yukarı kayarsa padding ekle */
    .block-container {
        padding-top: 1rem !important;
    }

    /* --- MEVCUT KODLARINIZ BURADAN DEVAM EDİYOR --- */
    
    /* Genel Uygulama Arka Planı */
    .stApp {
        background-color: #F8FAFC;
        font-family: 'Segoe UI', Roboto, Helvetica, Arial, sans-serif;
    }
    

<style>
    /* Genel Uygulama Arka Planı */
    .stApp {
        background-color: #F8FAFC;
        font-family: 'Segoe UI', Roboto, Helvetica, Arial, sans-serif;
    }
    
    /* Sidebar Tasarımı */
    [data-testid="stSidebar"] {
        background-color: #FFFFFF;
        border-right: 1px solid #E2E8F0;
    }
    
    /* Başlık Stilleri */
    h1, h2, h3, h4 {
        color: #0F172A;
        font-weight: 700;
        letter-spacing: -0.5px;
    }
    
    /* Dashboard Kartları (Metrikler) */
    div[data-testid="stMetric"] {
        background-color: #FFFFFF;
        padding: 24px;
        border-radius: 16px;
        border: 1px solid #E2E8F0;
        box-shadow: 0 4px 6px -1px rgba(0, 0, 0, 0.05);
        transition: all 0.3s ease;
    }
    div[data-testid="stMetric"]:hover {
        transform: translateY(-4px);
        box-shadow: 0 10px 15px -3px rgba(0, 0, 0, 0.1);
        border-color: #3B82F6;
    }
    div[data-testid="stMetric"] label {
        color: #64748B;
        font-size: 0.9rem;
        font-weight: 500;
    }
    div[data-testid="stMetric"] div[data-testid="stMetricValue"] {
        color: #0F172A;
        font-size: 1.8rem;
        font-weight: 700;
    }

    /* Özel Dashboard Karşılama Kartı */
    .dashboard-card {
        background-color: white;
        padding: 30px;
        border-radius: 12px;
        border: 1px solid #E2E8F0;
        box-shadow: 0 4px 6px rgba(0,0,0,0.02);
        margin-bottom: 20px;
    }
    
    /* Tablo (AgGrid) Çerçevesi */
    [data-testid="stDataFrame"] {
        border: 1px solid #E2E8F0;
        border-radius: 12px;
        overflow: hidden;
    }
    thead tr th {
        background-color: #F1F5F9 !important;
        color: #334155 !important;
        font-weight: 600 !important;
    }

    /* Buton Tasarımları */
    .stButton>button {
        border-radius: 8px;
        font-weight: 600;
        border: none;
        padding: 0.6rem 1.2rem;
        transition: all 0.2s ease;
        box-shadow: 0 2px 4px rgba(0,0,0,0.05);
    }
    /* Primary Buton (Lacivert/Mavi) */
    .stButton>button[kind="primary"] {
        background: linear-gradient(135deg, #1E3A8A 0%, #2563EB 100%);
        color: white;
        border: none;
    }
    .stButton>button[kind="primary"]:hover {
        box-shadow: 0 4px 12px rgba(37, 99, 235, 0.4);
        transform: scale(1.02);
    }
    /* Secondary Buton (Beyaz/Gri) */
    .stButton>button[kind="secondary"] {
        background-color: #FFFFFF;
        color: #475569;
        border: 1px solid #CBD5E1;
    }
    .stButton>button[kind="secondary"]:hover {
        background-color: #F8FAFC;
        border-color: #94A3B8;
        color: #0F172A;
    }
    
    /* Expander Başlığı */
    .streamlit-expanderHeader {
        background-color: white;
        border-radius: 8px;
        border: 1px solid #E2E8F0;
        font-weight: 600;
        color: #1E3A8A;
    }
</style>
""", unsafe_allow_html=True)

# ==============================================================================
# 5. SESSION STATE YÖNETİMİ
# ==============================================================================
# Uygulama boyunca verilerin korunması için state tanımlamaları
default_states = {
    'sayfa_secimi': "🏠 Ana Sayfa",
    'aktif_teklif_data': None,
    'islem_turu': "yeni",
    'secili_firma_adi': None,
    'tablo_verisi': None,
    'temp_img': None,
    'doviz_kurlari': {"TL": 1.0, "USD": 0.0, "EUR": 0.0, "GBP": 0.0},
    'tablo_aktif_para_birimi': "TL",
    'aktif_taslak_id': None,
    'aktif_detay_id': None,
    'aggrid_key': 0,
    'teklif_notlari': "",
    'genel_iskonto': 0.0,
    'nakliye_secimi': "HARİÇ",
    'montaj_secimi': "DAHİL",
    'form_proje': "",
    'form_no': "",
    'form_rev': "",
    'form_tarih': date.today(),
    'form_para': "TL",
    'form_kdv': 0,
    'secili_dil': "TR",
    'form_proje_kodu': "KOD",
    'sb_key': 0
}

for key, val in default_states.items():
    if key not in st.session_state:
        st.session_state[key] = val

# Global değişken (Sidebar ve ana ekran arası iletişim için)
secilen_firma_data = None 

# ==============================================================================
# 6. DİL SÖZLÜĞÜ (TR / EN)
# ==============================================================================
LABELS = {
    "TR": {
        "title1": "TEKLİF", 
        "title2": "PAKETİ",
        "client": "MÜŞTERİ ADI", 
        "project": "PROJE ADI",
        "date": "TARİH", 
        "no": "PROJE NO", 
        "rev": "REVİZYON NO",
        "thank_you": "Bizi tercih ettiğiniz için teşekkür ederiz.\nProjeniz ile ilgili hazırladığımız teklif dosyamız bilgilerinize sunulmuştur.\nTeklifimizin olumlu karşılanmasını umut eder, iyi çalışmalar dileriz.",
        "desc": "AÇIKLAMA", 
        "total": "TOPLAM",
        "discount": "İSKONTO", 
        "grand_total": "GENEL TOPLAM",
        "shipping": "NAKLİYE", 
        "assembly": "MONTAJ",
        "notes": "NOTLAR",
        "code": "KOD", 
        "image": "GÖRSEL", 
        "prod_name": "ÜRÜN ADI / DETAY",
        "dim": "ÖLÇÜ", 
        "qty": "MİK.", 
        "price": "FİYAT", 
        "disc_col": "İND.",
        "footer_factory": "BURSA FABRIKA - MERKEZ",
        "footer_ksa": "MEKKE OFİS",
        "footer_mk": "MAKEDONYA OFİS"
    },
    "EN": {
        "title1": "PROPOSAL", 
        "title2": "PACK",
        "client": "CLIENT NAME", 
        "project": "PROJECT NAME",
        "date": "DATE", 
        "no": "PROJECT NO", 
        "rev": "REVISION NO",
        "thank_you": "Thank you for choosing us.\nOur proposal file regarding your project is presented for your information.\nWe hope our offer meets your expectations and wish you a good day.",
        "desc": "DESCRIPTION", 
        "total": "TOTAL",
        "discount": "DISCOUNT", 
        "grand_total": "GRAND TOTAL",
        "shipping": "SHIPPING", 
        "assembly": "ASSEMBLY",
        "notes": "NOTES & CONDITIONS",
        "code": "CODE", 
        "image": "IMAGE", 
        "prod_name": "PRODUCT NAME / DETAIL",
        "dim": "DIM.", 
        "qty": "QTY", 
        "price": "PRICE", 
        "disc_col": "DISC.",
        "footer_factory": "BURSA FACTORY - HQ",
        "footer_ksa": "OFFICE SAUDI ARABIA",
        "footer_mk": "OFFICE MACEDONIA"
    }
}

# ==============================================================================
# 7. VERİTABANI İŞLEMLERİ
# ==============================================================================

def db_baglan():
    """Veritabanı bağlantısı oluşturur."""
    return sqlite3.connect("teklif_yonetim_sistemi.db")

def tablolari_olustur():
    """Gerekli tabloları oluşturur."""
    conn = db_baglan()
    c = conn.cursor()
    
    # Müşteriler Tablosu
    c.execute('''
        CREATE TABLE IF NOT EXISTS musteriler (
            id INTEGER PRIMARY KEY AUTOINCREMENT, 
            firma_adi TEXT, 
            yetkili_kisi TEXT, 
            adres TEXT
        )
    ''')
    
    # Teklifler Tablosu
    c.execute('''
        CREATE TABLE IF NOT EXISTS teklifler (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            musteri_id INTEGER,
            firma_adi TEXT,
            proje_adi TEXT,
            proje_no TEXT,
            revizyon TEXT,
            tarih TEXT,
            toplam_tutar REAL,
            urun_datasi TEXT,
            para_birimi TEXT,
            kdv_orani INTEGER,
            durum TEXT DEFAULT 'Yayında',
            ozel_notlar TEXT,
            genel_iskonto REAL DEFAULT 0,
            nakliye_durum TEXT DEFAULT 'HARİÇ',
            montaj_durum TEXT DEFAULT 'DAHİL',
            proje_kodu TEXT
        )
    ''')
    
    conn.commit()
    conn.close()

def sema_kontrol():
    """Veritabanı sütun eksikliklerini kontrol eder ve ekler."""
    conn = db_baglan()
    c = conn.cursor()
    
    cols = [
        "para_birimi", "kdv_orani", "durum", "ozel_notlar", 
        "genel_iskonto", "nakliye_durum", "montaj_durum", "proje_kodu"
    ]
    
    for col in cols:
        try: 
            c.execute(f"ALTER TABLE teklifler ADD COLUMN {col} TEXT")
        except: 
            pass
            
    conn.commit()
    conn.close()

def musterileri_getir():
    conn = db_baglan()
    df = pd.read_sql_query("SELECT * FROM musteriler", conn)
    conn.close()
    return df

def musteri_ekle(firma, yetkili, adres):
    conn = db_baglan()
    c = conn.cursor()
    c.execute(
        "INSERT INTO musteriler (firma_adi, yetkili_kisi, adres) VALUES (?, ?, ?)", 
        (firma, yetkili, adres)
    )
    conn.commit()
    conn.close()

def musteri_guncelle(id, yeni_firma, yeni_yetkili, yeni_adres):
    conn = db_baglan()
    c = conn.cursor()
    c.execute('''
        UPDATE musteriler 
        SET firma_adi=?, yetkili_kisi=?, adres=? 
        WHERE id=?
    ''', (yeni_firma, yeni_yetkili, yeni_adres, id))
    conn.commit()
    conn.close()

def musteri_sil(id):
    conn = db_baglan()
    c = conn.cursor()
    c.execute("DELETE FROM musteriler WHERE id = ?", (int(id),))
    conn.commit()
    conn.close()

def teklif_sil(id):
    conn = db_baglan()
    c = conn.cursor()
    c.execute("DELETE FROM teklifler WHERE id = ?", (id,))
    conn.commit()
    conn.close()

def teklif_ekle_veya_guncelle(id, musteri_id, firma_adi, proje_adi, proje_no, revizyon, tarih, toplam, urun_df, p_birim, kdv, durum="Yayında", notlar="", genel_iskonto=0, nakliye="HARİÇ", montaj="DAHİL", proje_kodu=""):
    """Teklif kaydetme veya güncelleme fonksiyonu."""
    conn = db_baglan()
    c = conn.cursor()
    
    # Gereksiz sütunları temizle ve JSON'a çevir
    kayit_df = urun_df.drop(columns=["GÖRSEL_DURUM", "DETAY_DURUM", "NO", "SİL", "GÖRSEL_GRID"], errors='ignore')
    urun_json = kayit_df.to_json(orient='records')
    
    try: 
        toplam_safe = float(toplam)
    except: 
        toplam_safe = 0.0
    
    yeni_id = id
    
    # Verileri hazırla
    veriler = (
        musteri_id, firma_adi, proje_adi, proje_no, revizyon, tarih, 
        toplam_safe, urun_json, p_birim, kdv, durum, notlar, 
        genel_iskonto, nakliye, montaj, proje_kodu
    )

    if id is None:
        # Yeni Kayıt
        c.execute('''
            INSERT INTO teklifler (
                musteri_id, firma_adi, proje_adi, proje_no, revizyon, tarih, 
                toplam_tutar, urun_datasi, para_birimi, kdv_orani, durum, 
                ozel_notlar, genel_iskonto, nakliye_durum, montaj_durum, proje_kodu
            ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
        ''', veriler)
        yeni_id = c.lastrowid
    else:
        # Güncelleme
        update_veriler = (
            proje_adi, proje_no, revizyon, tarih, toplam_safe, urun_json, 
            p_birim, kdv, durum, notlar, genel_iskonto, nakliye, montaj, proje_kodu, id
        )
        c.execute('''
            UPDATE teklifler SET 
                proje_adi=?, proje_no=?, revizyon=?, tarih=?, toplam_tutar=?, 
                urun_datasi=?, para_birimi=?, kdv_orani=?, durum=?, ozel_notlar=?, 
                genel_iskonto=?, nakliye_durum=?, montaj_durum=?, proje_kodu=?
            WHERE id=?
        ''', update_veriler)
        yeni_id = id
        
    conn.commit()
    conn.close()
    return yeni_id

def teklifleri_getir(musteri_id):
    conn = db_baglan()
    df = pd.read_sql_query("SELECT * FROM teklifler WHERE musteri_id = ? AND durum = 'Yayında'", conn, params=(musteri_id,))
    conn.close()
    return df

def taslaklari_getir(musteri_id):
    conn = db_baglan()
    df = pd.read_sql_query("SELECT * FROM teklifler WHERE musteri_id = ? AND durum = 'Taslak' ORDER BY id DESC", conn, params=(musteri_id,))
    conn.close()
    return df

def sonraki_revizyon(mevcut_rev):
    try:
        prefix = mevcut_rev[0] 
        num = int(mevcut_rev[1:]) 
        return f"{prefix}{num + 1:02d}" 
    except:
        return mevcut_rev + ".1" 

def base64_to_image(base64_string):
    try:
        if "," in base64_string:
            base64_string = base64_string.split(",")[1]
        img_data = base64.b64decode(base64_string)
        return Image.open(io.BytesIO(img_data))
    except:
        return None

# ==============================================================================
# 8. PDF VE EXCEL SINIFLARI (UNICODE HATASI DÜZELTİLMİŞ)
# ==============================================================================

class SaraksPDF(FPDF):
    def __init__(self, proje_adi, proje_no, rev_no, language="TR"):
        super().__init__()
        self.proje_adi = proje_adi
        self.proje_no = proje_no
        self.rev_no = rev_no
        self.language = language
        self.set_margins(3, 3, 3)
        self.set_auto_page_break(auto=True, margin=8)
        
        # --- KRİTİK FONT AYARI (UNICODE HATASINI ÇÖZEN KISIM) ---
        # FPDF varsayılan olarak Türkçe karakterleri (İ, ğ, ş) tanımaz.
        # Bu yüzden Windows sistemindeki Arial fontunu yüklüyoruz.
        font_dir = "C:\\Windows\\Fonts"
        regular_font = os.path.join(font_dir, "arial.ttf")
        bold_font = os.path.join(font_dir, "arialbd.ttf")
        
        self.font_ok = False
        
        if os.path.exists(regular_font):
            try:
                # uni=True parametresi UTF-8 desteği sağlar
                self.add_font("ArialTR", "", regular_font, uni=True)
                self.main_font = "ArialTR"
                self.font_ok = True
            except:
                self.main_font = "Arial"
        else:
            self.main_font = "Arial"

        if os.path.exists(bold_font) and self.font_ok:
            try:
                self.add_font("ArialTR", "B", bold_font, uni=True)
            except:
                pass

    def header(self):
        if self.page_no() == 2:
            if os.path.exists("logo.png"):
                try: self.image("logo.png", 3, 3, h=13)
                except: pass
            
            self.set_draw_color(18, 52, 86) 
            self.set_line_width(0.7) 
            self.line(3, 16, 207, 16)
            
            self.set_draw_color(0,0,0)
            self.set_line_width(0.2) 
            self.ln(5) 

    def footer(self):
        # Font hatası varsa footer metnini basitleştir
        if self.font_ok:
            l_dict = LABELS[self.language]
            fact_text = l_dict["footer_factory"]
        else:
            fact_text = "BURSA FABRIKA"

        # Sadece 2. sayfa ve sonrasında göster (veya isteğe bağlı olarak her sayfada)
        if self.page_no() == 2:
            # Sayfa sonundan 22 birim yukarı çık (Sabit başlangıç noktası)
            self.set_y(-22)
            
            # --- KRİTİK DÜZELTME: Y KOORDİNATINI SABİTLE ---
            # get_y() ile o anki Y değerini bir değişkene atıyoruz.
            # Artık tüm kutuları bu 'y_sabit' değerine göre çizeceğiz.
            y_sabit = self.get_y()
            
            self.set_font(self.main_font, "", 7)
            L = LABELS[self.language]
            
            self.set_draw_color(18, 52, 86)
            self.set_line_width(0.3)
            
            # KUTU YÜKSEKLİĞİ
            h_box = 18 

            # --- KUTU 1 (SOL) ---
            x1 = 3
            self.rect(x1, y_sabit, 68, h_box) # rect(x, y, w, h)
            
            # Kutu 1 Metinleri
            self.set_xy(x1 + 2, y_sabit + 2)
            self.set_font(self.main_font, "B", 7)
            self.cell(64, 4, fact_text, 0, 1, 'L')
            
            self.set_xy(x1 + 2, y_sabit + 6) # Y koordinatını elle ayarlıyoruz
            self.set_font(self.main_font, "", 7)
            self.multi_cell(64, 3.5, "Altinova mh. 3. Yazici sk. No:2/1\nOsmangazi / BURSA\nT: +90 224 215 20 99", 0, 'L')
            
            # --- KUTU 2 (ORTA) ---
            x2 = 3 + 68
            # BURADA ARTIK get_y() KULLANMIYORUZ, y_sabit KULLANIYORUZ
            self.rect(x2, y_sabit, 68, h_box) 
            
            # Kutu 2 Metinleri
            self.set_xy(x2, y_sabit + 2)
            self.set_font(self.main_font, "B", 7)
            self.cell(66, 4, L["footer_ksa"], 0, 1, 'R')
            
            self.set_xy(x2, y_sabit + 6)
            self.set_font(self.main_font, "", 7)
            self.multi_cell(66, 3.5, "Hajib Ibn Zararah 7011 Bani Muawiyah\nAl Madina Munawara\nT: +966 50 200 3603", 0, 'R')

            # --- KUTU 3 (SAĞ) ---
            x3 = x2 + 68
            # BURADA DA y_sabit KULLANIYORUZ
            self.rect(x3, y_sabit, 68, h_box)
            
            # Kutu 3 Metinleri
            self.set_xy(x3, y_sabit + 2)
            self.set_font(self.main_font, "B", 7)
            self.cell(68, 4, L["footer_mk"], 0, 1, 'C')
            
            self.set_xy(x3, y_sabit + 8) # Dikeyde ortalamak için biraz aşağı ittik
            self.set_font(self.main_font, "", 7)
            self.cell(68, 4, "COMING SOON - HACKOPO", 0, 1, 'C')

def create_pdf(firma_data, proje_data, df_urunler, ara_toplam, genel_iskonto_tutar, genel_toplam, para_birimi, kdv_orani, notlar_text, nakliye_durum, montaj_durum, language="TR"):
    L = LABELS[language]
    pdf = SaraksPDF(proje_data['adi'], proje_data['no'], proje_data['rev'], language)
    pdf.add_page()
    main_font = pdf.main_font
    
    # --- 1. KAPAK SAYFASI ---
    pdf.set_draw_color(18, 52, 86)
    pdf.set_line_width(6.0) 
    pdf.rect(5, 5, 200, 287) 
    
    pdf.set_xy(10, 13) 
    pdf.set_font(main_font, "B", 36) 
    pdf.set_text_color(18, 52, 86) 
    pdf.cell(0, 17, "FURNITURE", 0, 1, 'L')
    
    pdf.set_x(10)
    pdf.set_text_color(0, 0, 0) 
    max_text_width = pdf.get_string_width("INTERIOR FIT-OUT")
    pdf.cell(0, 17, "INTERIOR FIT-OUT", 0, 1, 'L')
    
    pdf.set_x(10)
    pdf.cell(0, 17, "CONTRACTING", 0, 1, 'L')
    
    pdf.set_draw_color(18, 52, 86)
    pdf.set_line_width(1.5)
    pdf.line(13, 65, 10 + 83, 65) 
    
    pdf.set_y(130)
    pdf.set_font(main_font, "", 55) 
    pdf.set_text_color(18, 52, 86) 
    pdf.cell(0, 20, L["title1"], 0, 1, 'C')
    pdf.cell(0, 20, L["title2"], 0, 1, 'C')
    
    pdf.ln(7)
    pdf.set_font(main_font, "B", 16) 
    pdf.set_text_color(0, 0, 0)
    pdf.cell(0, 10, proje_data['adi'].upper(), 0, 1, 'C')
    
    if os.path.exists("logo.png"):
        logo_width = 90
        try:
            img = Image.open("logo.png")
            aspect_ratio = img.height / img.width
            logo_height = logo_width * aspect_ratio
        except:
            logo_height = 30 
        
        page_width = 210
        x_pos = (page_width - logo_width) / 2
        target_y = 294 - 7 - logo_height
        pdf.image("logo.png", x=x_pos, y=target_y, w=logo_width)
    
    # --- 2. SAYFA (ÖZET) ---
    pdf.add_page()
    pdf.set_line_width(0.2)
    pdf.set_draw_color(0, 0, 0)
    pdf.set_font(main_font, "", 10)
    pdf.set_text_color(0, 0, 0) 
    
    pdf.ln(30)
    
    content_x = 23 
    content_w = 164 
    
    pdf.set_x(content_x)
    
    label_w = 40
    def print_kunye(label, value):
        start_x = pdf.get_x() 
        pdf.set_font(main_font, "B", 9)
        pdf.set_text_color(0,0,0) 
        pdf.cell(label_w, 6, label, 0, 0, 'L')
        pdf.set_font(main_font, "", 9)
        pdf.cell(5, 6, ":", 0, 0, 'C')
        pdf.cell(0, 6, value, 0, 1, 'L')
        pdf.set_x(start_x) 

    print_kunye(L["client"], str(firma_data['firma_adi']))
    print_kunye(L["project"], str(proje_data['adi']))
    print_kunye(L["date"], str(proje_data['tarih']))
    print_kunye(L["no"], str(proje_data['no']))
    print_kunye(L["rev"], str(proje_data['rev']))
    
    pdf.ln(20) 
    
    pdf.set_x(content_x)
    
    pdf.set_font(main_font, "", 10) 
    pdf.multi_cell(content_w, 5, L["thank_you"], 0, 'C')
    
    pdf.ln(20)
    
    show_summary = False
    if "MAHAL" in df_urunler.columns:
        cl = df_urunler.copy()
        cl["MAHAL"] = cl["MAHAL"].fillna("Genel").replace(["", " "], "Genel")
        
        # Benzersiz mahalleleri giriş sırasına göre al
        unique_mahals = cl['MAHAL'].unique()
        cl['MAHAL'] = pd.Categorical(cl['MAHAL'], categories=unique_mahals, ordered=True)
        
        mo = cl.groupby("MAHAL", observed=True)["TOPLAM FİYAT"].sum().reset_index()
        mo = mo[mo["MAHAL"] != "Genel"]
        if not mo.empty:
            show_summary = True

    if show_summary:
        col_total_w = 40 
        col_desc_w = content_w - col_total_w 
        
        pdf.set_x(content_x)
        pdf.set_font(main_font, "B", 10)
        pdf.set_text_color(0,0,0) 
        pdf.set_fill_color(240, 240, 240)
        pdf.set_line_width(0.2)
        
        pdf.cell(col_desc_w, 8, L["desc"], 1, 0, 'C', fill=True)
        pdf.cell(col_total_w, 8, L["total"], 1, 1, 'C', fill=True)
        
        pdf.set_font(main_font, "", 10)
        for idx, m_row in mo.iterrows():
            pdf.set_x(content_x)
            pdf.cell(col_desc_w, 8, f"  {str(m_row['MAHAL'])}", 1, 0, 'L')
            pdf.cell(col_total_w, 8, format_para(m_row['TOPLAM FİYAT'], para_birimi, 2) + "  ", 1, 1, 'R')
        
        label_w_bottom = 40
        val_w_bottom = col_total_w
        offset_x = content_x + content_w - (label_w_bottom + val_w_bottom)
        
        def draw_summary_row(label, value):
            pdf.set_x(offset_x)
            pdf.set_font(main_font, "B", 10)
            pdf.set_text_color(0,0,0)
            pdf.cell(label_w_bottom, 8, label, 1, 0, 'C') 
            pdf.set_font(main_font, "", 10)
            pdf.cell(val_w_bottom, 8, f"{value}  ", 1, 1, 'R')

        draw_summary_row(L["total"], format_para(ara_toplam, para_birimi, 2))
        
        if genel_iskonto_tutar > 0:
            draw_summary_row(L["discount"], "-" + format_para(genel_iskonto_tutar, para_birimi, 2))
            draw_summary_row(L["grand_total"], format_para(genel_toplam, para_birimi, 2))
            
        nakliye_val = nakliye_durum.upper()
        montaj_val = montaj_durum.upper()
        if language == "EN":
             nakliye_val = "INCLUDED" if nakliye_val == "DAHİL" else "EXCLUDED"
             montaj_val = "INCLUDED" if montaj_val == "DAHİL" else "EXCLUDED"

        draw_summary_row(L["shipping"], nakliye_val)
        draw_summary_row(L["assembly"], montaj_val)

        pdf.ln(15)
    
    pdf.set_x(content_x)
    pdf.set_font(main_font, "B", 10)
    pdf.set_text_color(0,0,0) 
    pdf.cell(0, 6, L["notes"], 0, 1, 'L')
    pdf.set_x(content_x)
    pdf.set_font(main_font, "", 9)
    pdf.multi_cell(content_w, 5, notlar_text)
    
    # --- 3. SAYFA (ÜRÜN LİSTESİ) ---
    pdf.add_page()
    pdf.set_margins(8, 8, 8) 
    pdf.set_y(8) 
    pdf.set_line_width(0.2)
    pdf.ln(1) 
    pdf.set_font(main_font, "B", 9)
    pdf.set_text_color(0, 0, 0) 
    
    has_discount = False
    if "İSKONTO" in df_urunler.columns and df_urunler["İSKONTO"].sum() > 0:
        has_discount = True

    current_page_w = 194 
    w_code = 16
    w_img = 30  
    w_dim = 9
    w_qty = 11
    w_price = 23
    w_total = 28
    w_disc = 11 if has_discount else 0
    
    fixed_w = w_code + w_img + w_dim + w_qty + w_price + w_total + w_disc
    w_name = current_page_w - fixed_w 

    x = 8 
    y = pdf.get_y()
    h = 8
    
    pdf.set_draw_color(18, 52, 86)
    pdf.set_line_width(0.5) 
    pdf.set_fill_color(220, 230, 241) 

    def draw_cell(w, txt, align='C', fill=True):
        pdf.cell(w, h, txt, 1, 0, align, fill)

    pdf.set_x(8)
    draw_cell(w_code, L["code"])
    draw_cell(w_img, L["image"])
    draw_cell(w_name, L["prod_name"], 'L')
    draw_cell(w_dim, L["dim"])
    draw_cell(w_qty, L["qty"])
    draw_cell(w_price, L["price"], 'R')
    if has_discount:
        draw_cell(w_disc, L["disc_col"])
    draw_cell(w_total, L["total"], 'R')
    pdf.ln(h) 
    
    pdf.set_line_width(0.2)
    pdf.set_draw_color(0, 0, 0)
    pdf.set_font(main_font, "", 8)
    
    df_urunler["MAHAL"] = df_urunler["MAHAL"].fillna("Genel").replace(["", " "], "Genel")
    
    current_mahal = None
    
    for index, row in df_urunler.iterrows():
        pdf.set_text_color(0,0,0)
        row_mahal = row.get("MAHAL", "Genel")
        
        if row_mahal != current_mahal:
            if row_mahal != "Genel": 
                mahal_total = df_urunler[df_urunler["MAHAL"] == row_mahal]["TOPLAM FİYAT"].sum()
                pdf.ln(2)
                pdf.set_x(8)
                pdf.set_font(main_font, "B", 9)
                pdf.set_fill_color(220, 230, 241)
                
                pdf.cell(current_page_w - w_total, 8, f"  {row_mahal}", 1, 0, 'L', fill=True)
                pdf.cell(w_total, 8, format_para(mahal_total, para_birimi, 2), 1, 1, 'R', fill=True)
                pdf.set_font(main_font, "", 8)
            current_mahal = row_mahal

        line_height = 5
        desc_text = f"{row['ÜRÜN ADI']}\n{row['AÇIKLAMA']}"
        desc_lines = pdf.multi_cell(w_name, line_height, desc_text, split_only=True)
        num_lines = len(desc_lines)
        row_height = max(25, num_lines * line_height)
        
        # Görsel varsa yükseklik ayarı
        if row['GÖRSEL'] and len(str(row['GÖRSEL'])) > 20:
             row_height = max(row_height, 35)

        if pdf.get_y() + row_height > 280:
            pdf.add_page()
            pdf.set_margins(8, 8, 8) 
            pdf.set_line_width(0.2)
            pdf.set_x(8)
            pdf.set_y(8) 
            
            pdf.set_font(main_font, "B", 9)
            pdf.set_fill_color(220, 230, 241)
            pdf.set_draw_color(18, 52, 86)
            
            draw_cell(w_code, L["code"])
            draw_cell(w_img, L["image"])
            draw_cell(w_name, L["prod_name"], 'L')
            draw_cell(w_dim, L["dim"])
            draw_cell(w_qty, L["qty"])
            draw_cell(w_price, L["price"], 'R')
            if has_discount:
                draw_cell(w_disc, L["disc_col"])
            draw_cell(w_total, L["total"], 'R')
            pdf.ln(h)
            
            pdf.set_font(main_font, "", 8)
            pdf.set_line_width(0.2)
            pdf.set_draw_color(0,0,0)

        y_start = pdf.get_y()
        x_curr = 8 
        
        # KOD
        pdf.set_xy(x_curr, y_start + (row_height/2) - 3)
        pdf.cell(w_code, 6, str(row['KOD']), 0, 0, 'C') 
        pdf.rect(x_curr, y_start, w_code, row_height) 
        x_curr += w_code
        
        # GÖRSEL
        pdf.rect(x_curr, y_start, w_img, row_height) 
        if row['GÖRSEL'] and len(str(row['GÖRSEL'])) > 20:
            try:
                img_data = base64.b64decode(row['GÖRSEL'].split(",")[1])
                with tempfile.NamedTemporaryFile(delete=False, suffix=".jpg") as tmp_file:
                    tmp_file.write(img_data)
                    tmp_path = tmp_file.name
                
                with Image.open(tmp_path) as pil_img:
                    orig_w, orig_h = pil_img.size
                    max_w = w_img - 2
                    max_h = row_height - 2
                    ratio = min(max_w / orig_w, max_h / orig_h)
                    new_w = orig_w * ratio
                    new_h = orig_h * ratio
                    x_centered = x_curr + (w_img - new_w) / 2
                    y_centered = y_start + (row_height - new_h) / 2
                    pdf.image(tmp_path, x=x_centered, y=y_centered, w=new_w, h=new_h)
                os.unlink(tmp_path)
            except:
                pass
        x_curr += w_img

        # ÜRÜN ADI
        pdf.set_xy(x_curr, y_start)
        pdf.multi_cell(w_name, line_height, desc_text, 0, 'L')
        pdf.rect(x_curr, y_start, w_name, row_height)
        x_curr += w_name
        
        # ÖLÇÜ
        pdf.rect(x_curr, y_start, w_dim, row_height)
        center_x = x_curr + w_dim / 2
        center_y = y_start + row_height / 2
        text_val = str(row['ÖLÇÜ'])
        text_w = pdf.get_string_width(text_val)
        
        try:
            with pdf.rotation(90, x=center_x, y=center_y):
                 pdf.text(center_x - (text_w / 2), center_y + 1.5, text_val)
        except:
             pdf.text(center_x - (text_w / 2), center_y + 1.5, text_val)
             
        x_curr += w_dim
        
        # MİKTAR
        pdf.rect(x_curr, y_start, w_qty, row_height)
        pdf.set_xy(x_curr, y_start + (row_height/2) - 4)
        
        miktar_degeri = row['MİKTAR']
        try:
            miktar_f = float(miktar_degeri)
            if miktar_f.is_integer():
                miktar_str = str(int(miktar_f))
            else:
                miktar_str = str(miktar_f)
        except:
            miktar_str = str(miktar_degeri)

        pdf.multi_cell(w_qty, 4, f"{miktar_str}\n{row['BİRİM']}", 0, 'C')
        x_curr += w_qty
        
        # BİRİM FİYAT
        pdf.set_xy(x_curr, y_start + (row_height/2) - 3)
        pdf.cell(w_price, 6, format_para(row['BİRİM FİYAT'], "", 1), 0, 0, 'R')
        pdf.rect(x_curr, y_start, w_price, row_height)
        x_curr += w_price
        
        # İSKONTO
        if has_discount:
            disc_val = row.get("İSKONTO", 0)
            disc_txt = f"%{disc_val}" if disc_val > 0 else "-"
            pdf.set_xy(x_curr, y_start + (row_height/2) - 3)
            pdf.cell(w_disc, 6, disc_txt, 0, 0, 'C')
            pdf.rect(x_curr, y_start, w_disc, row_height)
            x_curr += w_disc
        
        # TOPLAM FİYAT
        pdf.set_xy(x_curr, y_start + (row_height/2) - 3)
        pdf.cell(w_total, 6, format_para(row['TOPLAM FİYAT'], "", 2), 0, 0, 'R')
        pdf.rect(x_curr, y_start, w_total, row_height)
        
        pdf.set_xy(8, y_start + row_height)
    
    pdf.set_font(main_font, "B", 10)
    pdf.set_fill_color(240, 240, 240)
    
    label_w = current_page_w - w_total
    
    # Eğer İskonto varsa Detaylı Göster (TOPLAM / İSKONTO / GENEL TOPLAM)
    if genel_iskonto_tutar > 0:
        # 1. Satır: Ara Toplam (TOPLAM)
        pdf.cell(label_w, 6, L["total"], 1, 0, 'R', fill=True)
        pdf.cell(w_total, 6, format_para(ara_toplam, para_birimi, 2), 1, 1, 'R', fill=True)
        
        # 2. Satır: İskonto Tutarı
        pdf.set_x(8)
        pdf.cell(label_w, 6, L["discount"], 1, 0, 'R', fill=True)
        pdf.cell(w_total, 6, "-" + format_para(genel_iskonto_tutar, para_birimi, 2), 1, 1, 'R', fill=True)
        
        # 3. Satır: Genel Toplam
        pdf.set_x(8)
        pdf.cell(label_w, 8, L["grand_total"], 1, 0, 'R', fill=True)
        pdf.cell(w_total, 8, format_para(genel_toplam, para_birimi, 2), 1, 1, 'R', fill=True)
        
    else:
        # İskonto yoksa sadece Genel Toplam göster (Eski hali)
        pdf.cell(label_w, 10, L["grand_total"], 1, 0, 'R', fill=True)
        pdf.cell(w_total, 10, format_para(genel_toplam, para_birimi, 2), 1, 1, 'R', fill=True)
    
    return pdf.output(dest='S').encode('latin-1')

# ==============================================================================
# SÖZLEŞME MOTORU (TARİH FORMATI DÜZELTİLMİŞ - TAM VERSİYON)
# ==============================================================================
def create_contract_pdf(
    # Sabit Veriler
    proje_adi, toplam_tutar, para_birimi,
    # Düzenlenebilir Müşteri Verileri
    mus_adi, mus_adres, mus_vd, 
    # Düzenlenebilir Sözleşme Verileri
    sozlesme_tarihi, bitis_tarihi, sehir,
    gecikme_orani, fesih_gun, garanti_suresi, odeme_plani
):
    
    pdf = FPDF()
    pdf.add_page()
    
    # Font Ayarı
    font_dir = os.path.join(os.environ.get('WINDIR', 'C:\\Windows'), 'Fonts')
    if os.path.exists(os.path.join(font_dir, 'arial.ttf')):
        pdf.add_font('ArialTR', '', os.path.join(font_dir, 'arial.ttf'), uni=True)
        pdf.add_font('ArialTR', 'B', os.path.join(font_dir, 'arialbd.ttf'), uni=True)
        main_font = 'ArialTR'
    else:
        main_font = 'Arial'

    # --- TARİH FORMATLAMA (DÜZELTME BURADA YAPILDI) ---
    # Gelen tarihi (2026-02-16) alıp (16.02.2026) formatına çeviriyoruz
    try:
        s_tarih_str = pd.to_datetime(str(sozlesme_tarihi)).strftime('%d.%m.%Y')
        b_tarih_str = pd.to_datetime(str(bitis_tarihi)).strftime('%d.%m.%Y')
    except:
        # Eğer hata olursa olduğu gibi yazsın
        s_tarih_str = str(sozlesme_tarihi)
        b_tarih_str = str(bitis_tarihi)

    # --- YÜKLENİCİ BİLGİLERİ (SABİT) ---
    YUKLENICI_UNVAN = "ONUR ALIŞIK – SARAKS MOBİLYA"
    YUKLENICI_ADRES = "ALTINOVA MH. 3.YAZICI SK. NO:2 OSMANGAZİ - BURSA"
    YUKLENICI_VD = "27385398522 - ULUDAĞ"

    # 1. BAŞLIK
    pdf.set_font(main_font, 'B', 14)
    pdf.cell(0, 10, "YÜKLENİCİ HİZMET SÖZLEŞMESİ", 0, 1, 'C')
    pdf.ln(5)

    # 2. GİRİŞ BİLGİLERİ (TABLO)
    def satir_yaz(baslik, deger):
        pdf.set_font(main_font, 'B', 9)
        pdf.cell(55, 5, baslik, 0, 0, 'L')
        pdf.set_font(main_font, '', 9)
        pdf.cell(3, 5, ":", 0, 0, 'C')
        pdf.multi_cell(0, 5, str(deger))

    # Tutar Formatı ve YAZIYLA YAZILMASI
    try:
        tutar_rakam = format_para(toplam_tutar, para_birimi)
        tutar_yazi = sayiyi_yaziya_cevir(toplam_tutar)
        pb_yazi = "TÜRK LİRASI" if "TL" in para_birimi or "TRY" in para_birimi else para_birimi
        tutar_komple = f"{tutar_rakam} ( {tutar_yazi} {pb_yazi} )"
    except:
        tutar_komple = f"{toplam_tutar} {para_birimi}"

    satir_yaz("SÖZLEŞMENİN KONUSU", f"{proje_adi} işinin yapılması")
    satir_yaz("YÜKLENİCİ FİRMANIN ADI", YUKLENICI_UNVAN)
    satir_yaz("YÜKLENİCİ FİRMA ADRESİ", YUKLENICI_ADRES)
    satir_yaz("VERGİ DAİRESİ VE NUMARASI", YUKLENICI_VD)
    
    # Müşteri Bilgileri
    satir_yaz("İŞVEREN ADI VEYA ÜNVANI", mus_adi)
    satir_yaz("İŞVEREN ADRESİ", mus_adres)
    satir_yaz("İŞVEREN VD / T.C. KİMLİK NO", mus_vd)

    satir_yaz("SÖZLEŞME BEDELİ", tutar_komple)
    satir_yaz("SÖZLEŞME TARİHİ", s_tarih_str) # Düzeltilmiş tarih
    satir_yaz("İŞİN BİTİM TARİHİ", b_tarih_str) # Düzeltilmiş tarih
    
    pdf.ln(5)

    # --- MADDELER ---
    def madde_baslik(txt):
        pdf.ln(3)
        pdf.set_font(main_font, 'B', 9)
        pdf.cell(0, 5, txt, 0, 1, 'L')
        pdf.set_font(main_font, '', 8)

    def madde_icerik(txt):
        pdf.multi_cell(0, 4, txt)

    # Madde 1
    madde_baslik("Madde 1- Taraflar:")
    madde_icerik(f"Bir tarafta SARAKS MOBİLYA adına hareket eden ONUR ALIŞIK ile diğer tarafta {mus_adi} arasında aşağıdaki şartlar dahilinde bu sözleşme akdedilmiştir. Sözleşme metninde SARAKS MOBİLYA-ONUR ALIŞIK (YÜKLENİCİ), {mus_adi} (İŞVEREN), kelimeleri ile ifade edilmiştir.")

    # Madde 2
    madde_baslik("Madde 2- Sözleşmenin Konusu:")
    madde_icerik(f"İŞVEREN tarafından yaptırılacak olan: {proje_adi} projesine ait işin yaptırılmasıdır. İş Anahtar teslimi götürü fiyat esasına göre yapılacak olup işin götürü fiyat tutarı {tutar_komple}'dir. YÜKLENİCİ söz konusu işi özel şartnamedeki hususlara göre Ana Sözleşme ve Sözleşme şartlarına uygun olarak zamanında yapmayı kabul ve taahhüt etmiştir.")

    # Madde 3
    madde_baslik("Madde 3- Sözleşmenin Ekleri:")
    madde_icerik("1- Proje Teklif Dosyası\n3- Proje Uygulama Çizimleri")

    # Madde 4 (TAM METİN)
    madde_baslik("Madde 4- Müddet ve Gecikme Cezası:")
    text_m4 = (
        f"İşin müddeti “Mücbir Sebepler” dışında iş bu müddet dahilinde bitmediği takdirde geçecek beher gün için "
        f"YÜKLENİCİ’den Sözleşme Tutarının {gecikme_orani} Oranında gecikme cezası kesilecektir. "
        f"Şu kadar ki; gecikme {fesih_gun} günü geçerse İŞVEREN firma gecikme cezasını almaya devam ederek beklemeye "
        f"veya mahkeme kanalı ile tespit yaptırmaya, ayrıca bir karar almaya, protesto çekmeye lüzum kalmaksızın "
        f"YÜKLENİCİ’nin nam ve hesabına işi bir başkasına yaptırmaya veya sözleşmeyi fesih ederek, uğradığı zararları "
        f"veya üçüncü şahsa ödemeye mecbur kalacağı cezai şart ya da tazminatı YÜKLENİCİ’den talep etmeye yetkilidir."
    )
    madde_icerik(text_m4)

    # Madde 5 (TAM METİN)
    madde_baslik("Madde 5- İş Programı:")
    text_m5 = (
        f"YÜKLENİCİ, sözleşmenin imzalanmasına müteakip 3 gün içerisinde taahhüdünün 5.maddesinde yazılı süre içerisinde "
        f"mevcut işi bitireceğini belirten bir iş programını İŞVEREN firmaya vermekle mükelleftir. "
        f"Bu iş programı işveren firmaya verilmemesi halinde gecikilen her gün için Sözleşme Tutarının {gecikme_orani} Oranında "
        f"cezayı işverene ödemekle yükümlüdür. İşveren, YÜKLENİCİ firma tarafından yapılmış iş programına işi durumunu göre "
        f"müdahil olarak ara temrinler tespit edebilir. YÜKLENİCİ bu tespitlere itiraz edemez."
    )
    madde_icerik(text_m5)

    # Madde 6 (TAM METİN)
    madde_baslik("Madde 6- Vergi ve Vesair Masraflar:")
    madde_icerik("Taahhüdün ifasına ait her türlü vergi, resim ve harçlar ile Noter masrafları her çeşit sigorta primleri, işçi ve işveren hisseleri, İşsizlik Sigortası Primleri, fazla mesai, ikramiye ve Pazar yevmiyeleri gibi iş kanununun gerektirdiği bilcümle vecibeler YÜKLENİCİ’a aittir. YÜKLENİCİ mevcut vergi, resim, prim ve harçların artması veya rayiçlerin yükselmesi veya yeniden vergi, prim ve harçlar ihdası gibi sebeplere dayanarak yeni fiyat veya süre uzatılması gibi taleplerde bulunamaz.")

    # Madde 7
    madde_baslik("Madde 7- İmalatın ve İhzaratın Muhafazası:")
    madde_icerik("YÜKLENİCİun kendi işyerinde veya inşaat mahallinde yaptığı ihzarat, İŞVEREN firmasının teslim ettiği malzemeler ve imalatların İŞVEREN firmasına teslimine kadar muhafazası ve mesuliyeti YÜKLENİCİ’a aittir.")

    # Madde 8
    madde_baslik("Madde 8- Kusurlu Malzeme, İmalat ve Hasarlar:")
    madde_icerik("Şartnamelerde yazılı hükümlere uymayan veya fen gereği olarak belli vasıfları ve şartları haiz olmayan malzeme ile proje ve şartnamesine ve tekniğine uymayan imalat red olunur. YÜKLENİCİ kabul edilmeyen malzemeyi değiştirmeye veya düzeltmeye veya yeniden imal etmeye mecburdur. Bu yüzden hasıl olabilecek gecikmeler iş müddetinin ve ara terminlerin uzatılmasını gerektirmez. İŞVEREN firması, bu gibi kusurlu malzeme ve imalattan mütevellit maruz kalacağı ziyanı ayrıca YÜKLENİCİ’den talep etmek hakkına haizdir.")

    # Madde 9
    madde_baslik("Madde 9- İmalat Miktarının Artma veya Eksilmesi:")
    madde_icerik("YÜKLENİCİ’nin Kapsamı dışında ve karşılıklı fiyat mutabakatı sağlanması kaydı ile İŞVEREN firmasının göreceği lüzum üzerine fazla veya eksik iş yaptırmaya yetkilidir. Ancak iş miktarındaki artma veya eksilme yekûn bedelin % 30 ’dan fazla veya eksik olamaz. YÜKLENİCİ işlerin bu nispet dahilinde ki artma veya eksilmeden dolayı zarar ve ziyan gibi bir talep ve itirazda bulunamaz.")

    # Madde 10
    madde_baslik("Madde 10- Taahhüdün Devri:")
    madde_icerik("YÜKLENİCİ bu sözleşme ile taahhüt ettiği işleri İŞVEREN firmasının yazılı onayını almaksızın kısmen veya tamamen başka birine devir ve temlik edemez.")

    # Madde 11
    madde_baslik("Madde 11- İmalatın Teslimi:")
    madde_icerik("YÜKLENİCİ tarafından işin ikmal edildiği İŞVEREN firmasına bildirilmesini müteakip İŞVEREN firmasınca teşkil olunacak heyet marifetiyle imalatın monte edildiği mahalde işin kontrolu yapılıp, kabul edilerek veya 9.madde uyarınca işlem yapılarak bir tutanak tanzim edilecektir. YÜKLENİCİ’nin imalat yerinden İŞVEREN firmanın iş yerine kadar olan her türlü nakliye, yükleme, boşaltma, istifleme giderleri ile bu meyanda meydana gelecek her türlü hasar ve kusurlar YÜKLENİCİ’ye aittir.")

    # Madde 12
    madde_baslik("Madde 12- Garanti Müddeti:")
    madde_icerik(f"YÜKLENİCİ firma tarafından yapılan işlerin garanti süresi; aksi belirtilmediği sürece {garanti_suresi} olarak kabul edilir.")

    # Madde 13
    madde_baslik("Madde 13- Ödeme Şekli:")
    madde_icerik(f"{odeme_plani}")

    # Madde 14 (TAM METİN)
    madde_baslik("Madde 14- Emniyet Tedbirleri:")
    text_m14 = (
        "Herhangi bir ameliyenin yapılması dolayısı ile vukua gelebilecek kazalardan korunmak için YÜKLENİCİ iş güvenliği "
        "ve iş tüzüğü ile tespit edilen bütün tedbirleri alacak ve kazalardan korunma usul ve çarelerinin işçi ve personeline "
        "öğretecektir. YÜKLENİCİ, kazalara karşı her türlü emniyet tedbirlerini almakla mükellef olup, gerek ihmal, "
        "dikkatsizlik veya tedbirsizlikten, gerekse ehliyetsiz işçi kullanmaktan veya herhangi bir başka sebeplerle vuku "
        "bulacak kazalardan mesul olup, kazaya uğrayacak işçi, personel ve üçüncü kişilerin tedavi ve kendilerine, ailelerine "
        "verilecek tazminat, mahkeme masrafları ve sair masrafları tamamen YÜKLENİCİ’a aittir. İŞVEREN firması bu nedenle "
        "herhangi bir talep ve/veya ödemeye maruz kalırsa bu bedel YÜKLENİCİ tarafından aynen karşılanacaktır."
    )
    madde_icerik(text_m14)

    # Madde 15
    madde_baslik("Madde 15- Kanuni İkâmetgah:")
    madde_icerik(f"YÜKLENİCİ {YUKLENICI_ADRES} adresini kanuni ikametgâh olarak göstermiş olup, bu adrese yapılacak tebligat aynı günde YÜKLENİCİ’un kendisine yapılmış sayılır.")

    # Madde 16
    madde_baslik("Madde 16- İhtilafların Halli:")
    madde_icerik("Bu sözleşmenin tatbikinden doğacak her türlü ihtilafların halli, mercii T.C. Bursa Mahkemeleri ve İcra Daireleridir.")

    # Madde 17
    madde_baslik("Madde 17- Sözleşme Ekleri")
    madde_icerik(f"Sözleşme eki olarak madde 3 de yer alan ekler sözleşmenin tamamlayıcısı olup, ayrılmaz bir parçasıdır.")
    
    # Madde 18
    madde_baslik("Madde 18- Tarih ve Yer:")
    madde_icerik(f"Bu sözleşme taraflar arasında {s_tarih_str} tarihinde {sehir}'da tanzim ve imza edilmiştir.")
    
    # 3. İMZA ALANI
    pdf.ln(10)
    # Sayfa sonu kontrolü
    if pdf.get_y() > 240: pdf.add_page()
    
    pdf.set_font(main_font, 'B', 10)
    pdf.cell(95, 6, "İŞVEREN FİRMA", 0, 0, 'C')
    pdf.cell(95, 6, "YÜKLENİCİ FİRMA", 0, 1, 'C')
    
    pdf.ln(2)
    pdf.set_font(main_font, 'B', 9)
    # Müşteri adı
    pdf.cell(95, 6, str(mus_adi)[:45], 0, 0, 'C') 
    pdf.cell(95, 6, "SARAKS MOBİLYA - ONUR ALIŞIK", 0, 1, 'C')
    
    pdf.ln(5)
    pdf.set_font(main_font, '', 8)
    pdf.cell(95, 6, "İŞVEREN ADINA", 0, 0, 'C')
    pdf.cell(95, 6, "YÜKLENİCİ ADINA", 0, 1, 'C')

    pdf.ln(5)
    pdf.set_font(main_font, 'B', 9)
    pdf.cell(95, 6, "", 0, 0, 'C')
    pdf.cell(95, 6, "ONUR ALIŞIK", 0, 1, 'C')

    return pdf.output(dest='S').encode('latin-1')

# --- EXCEL EXPORT FONKSİYONU (MAHAL GRUPLU & DİNAMİK İSKONTO) ---
# Gerekli import (Kodun en başına ekleyin)
from xlsxwriter.utility import xl_rowcol_to_cell

def create_excel(firma_data, proje_data, df_urunler, ara_toplam, genel_iskonto_tutar, genel_toplam, para_birimi, notlar_text, nakliye_durum, montaj_durum):
    output = io.BytesIO()
    workbook = xlsxwriter.Workbook(output, {'in_memory': True})
    worksheet = workbook.add_worksheet("Teklif Detayı")
    
    # --- 1. FORMATLAR ---
    fmt_baslik_proje = workbook.add_format({'bold': True, 'font_size': 12, 'align': 'left', 'font_color': '#1E3A8A'})
    fmt_baslik_tablo = workbook.add_format({'bold': True, 'align': 'center', 'valign': 'vcenter', 'bg_color': '#1E3A8A', 'font_color': 'white', 'border': 1, 'text_wrap': True})
    
    # Standart Hücre Formatları
    fmt_center = workbook.add_format({'align': 'center', 'valign': 'vcenter', 'border': 1, 'text_wrap': True})
    fmt_left = workbook.add_format({'align': 'left', 'valign': 'vcenter', 'border': 1, 'text_wrap': True})
    fmt_money = workbook.add_format({'num_format': '#,##0.00', 'align': 'right', 'valign': 'vcenter', 'border': 1})
    
    # YENİ: Dikey Yazı Formatı (Ölçü İçin)
    fmt_dikey = workbook.add_format({
        'align': 'center', 
        'valign': 'vcenter', 
        'border': 1, 
        'text_wrap': True,
        'rotation': 90  # Metni 90 derece döndür
    })

    # Mahal Ara Toplam
    fmt_mahal_toplam_label = workbook.add_format({'bold': True, 'align': 'right', 'bg_color': '#FFF7E6', 'border': 1, 'font_color': '#B45309'})
    fmt_mahal_toplam_val = workbook.add_format({'bold': True, 'num_format': '#,##0.00', 'align': 'right', 'bg_color': '#FFF7E6', 'border': 1, 'font_color': '#B45309'})

    # Genel Toplam
    fmt_total_label = workbook.add_format({'bold': True, 'align': 'right', 'bg_color': '#F1F5F9', 'border': 1})
    fmt_total_val = workbook.add_format({'bold': True, 'num_format': '#,##0.00', 'align': 'right', 'bg_color': '#F1F5F9', 'border': 1})
    
    # --- 2. PROJE BİLGİLERİ ---
    worksheet.merge_range('A1:H1', f"MÜŞTERİ: {firma_data['firma_adi']}", fmt_baslik_proje)
    worksheet.write('A2', f"Proje Adı: {proje_data['adi']}")
    worksheet.write('A3', f"Proje No: {proje_data['no']}")
    worksheet.write('F2', f"Tarih: {str(proje_data['tarih'])}")
    worksheet.write('F3', f"Revizyon: {proje_data.get('rev', '-')}")
    
    # --- 3. VERİ HAZIRLIĞI ---
    df_export = df_urunler.copy()
    ignore_cols = ["GÖRSEL_GRID", "GÖRSEL_DURUM", "DETAY_DURUM", "SİL", "NO"]
    df_export = df_export.drop(columns=[c for c in ignore_cols if c in df_export.columns], errors='ignore')
    df_export = temizle_ve_sayiya_cevir(df_export, ["BİRİM FİYAT", "MİKTAR", "İSKONTO", "GİZLİ_İSKONTO"])

    # GİZLİ İSKONTO
    if "GİZLİ_İSKONTO" in df_export.columns:
        df_export["BİRİM FİYAT"] = df_export["BİRİM FİYAT"] * (1 - df_export["GİZLİ_İSKONTO"]/100)
        df_export["TOPLAM FİYAT"] = df_export["MİKTAR"] * df_export["BİRİM FİYAT"] * (1 - df_export["İSKONTO"]/100)
        df_export = df_export.drop(columns=["GİZLİ_İSKONTO"], errors='ignore')

    # GÖRÜNÜR İSKONTO KONTROLÜ
    iskonto_var = False
    if "İSKONTO" in df_export.columns and df_export["İSKONTO"].sum() > 0:
        iskonto_var = True
    
    # --- 4. SÜTUN HARİTASI ---
    columns_map = {} 
    col_idx = 0
    columns_map[col_idx] = ("KOD", 6); col_kod = col_idx; col_idx += 1
    columns_map[col_idx] = ("GÖRSEL", 25); col_img = col_idx; col_idx += 1
    columns_map[col_idx] = ("MAHAL", 8); col_mahal = col_idx; col_idx += 1
    columns_map[col_idx] = ("ÜRÜN ADI", 15); col_ad = col_idx; col_idx += 1
    columns_map[col_idx] = ("AÇIKLAMA", 30); col_aciklama = col_idx; col_idx += 1
    columns_map[col_idx] = ("ÖLÇÜ", 6); col_olcu = col_idx; col_idx += 1 # Genişliği biraz daralttık çünkü dik yazılacak
    columns_map[col_idx] = ("MİKTAR", 5); col_miktar = col_idx; col_idx += 1
    columns_map[col_idx] = ("BİRİM", 5); col_birim = col_idx; col_idx += 1
    columns_map[col_idx] = (f"BİRİM FİYAT ({para_birimi})", 10); col_fiyat = col_idx; col_idx += 1
    
    col_isk = -1
    if iskonto_var:
        columns_map[col_idx] = ("İSK. (%)", 10); col_isk = col_idx; col_idx += 1
        
    columns_map[col_idx] = (f"TOPLAM ({para_birimi})", 15); col_toplam = col_idx
    
    # Başlıkları Yaz
    row_idx = 5
    for c_idx, (name, width) in columns_map.items():
        worksheet.write(row_idx, c_idx, name, fmt_baslik_tablo)
        worksheet.set_column(c_idx, c_idx, width)

    # --- 5. SATIRLARI YAZDIRMA ---
    row_idx += 1
    unique_mahals = df_export['MAHAL'].unique()
    
    row_height = 100
    img_target_h = 120
    img_target_w = 160
    
    mahal_ara_toplam_hucreleri = []

    for mahal in unique_mahals:
        mahal_df = df_export[df_export['MAHAL'] == mahal]
        mahal_start_row = row_idx 
        
        for _, row in mahal_df.iterrows():
            worksheet.set_row(row_idx, row_height)
            
            # Verileri Yaz
            worksheet.write(row_idx, col_kod, row.get("KOD", ""), fmt_center)
            
            # GÖRSEL HÜCRESİ (Boş haliyle de ortalı)
            worksheet.write(row_idx, col_img, "", fmt_center)
            
            worksheet.write(row_idx, col_mahal, row.get("MAHAL", ""), fmt_center)
            
            # ÜRÜN ADI: Sola dayalı (fmt_left) yerine Ortalı (fmt_center) yapıldı
            worksheet.write(row_idx, col_ad, row.get("ÜRÜN ADI", ""), fmt_center)
            
            # Açıklama Sola Dayalı kalmalı (okunabilirlik için)
            worksheet.write(row_idx, col_aciklama, row.get("AÇIKLAMA", ""), fmt_left)
            
            # ÖLÇÜ: Yeni Dikey Format
            worksheet.write(row_idx, col_olcu, row.get("ÖLÇÜ", ""), fmt_dikey)
            
            worksheet.write(row_idx, col_miktar, row.get("MİKTAR", 0), fmt_center)
            worksheet.write(row_idx, col_birim, row.get("BİRİM", ""), fmt_center)
            worksheet.write(row_idx, col_fiyat, row.get("BİRİM FİYAT", 0), fmt_money)
            
            if iskonto_var:
                worksheet.write(row_idx, col_isk, row.get("İSKONTO", 0), fmt_center)
            
            # FORMÜL
            cell_miktar = xl_rowcol_to_cell(row_idx, col_miktar)
            cell_fiyat = xl_rowcol_to_cell(row_idx, col_fiyat)
            
            if iskonto_var:
                cell_isk = xl_rowcol_to_cell(row_idx, col_isk)
                formula = f"={cell_miktar}*{cell_fiyat}*(1-{cell_isk}/100)"
            else:
                formula = f"={cell_miktar}*{cell_fiyat}"
            
            worksheet.write_formula(row_idx, col_toplam, formula, fmt_money, row.get("TOPLAM FİYAT", 0))
            
            # --- GÖRSEL İŞLEME (ORTALAMA) ---
            gorsel_b64 = row.get("GÖRSEL", "")
            if gorsel_b64 and len(str(gorsel_b64)) > 50:
                try:
                    if "," in gorsel_b64: gorsel_b64 = gorsel_b64.split(",")[1]
                    img_data = base64.b64decode(gorsel_b64)
                    with Image.open(io.BytesIO(img_data)) as img:
                        img.thumbnail((img_target_w, img_target_h), Image.Resampling.LANCZOS)
                        img_byte_arr = io.BytesIO()
                        img.save(img_byte_arr, format='PNG')
                        img_byte_arr.seek(0)
                        
                        # Hücre Boyutları (Yaklaşık Piksel)
                        # Excel Column Width 25 ~= 180px
                        # Excel Row Height 100 ~= 133px
                        cell_w_px, cell_h_px = 180, 133
                        
                        # Ortalamak için offset hesabı
                        x_off = max(2, (cell_w_px - img.width) / 2)
                        y_off = max(2, (cell_h_px - img.height) / 2)
                        
                        worksheet.insert_image(row_idx, col_img, "img.png", {
                            'image_data': img_byte_arr,
                            'x_offset': x_off,
                            'y_offset': y_off,
                            'object_position': 1
                        })
                except: pass
            # ---------------------
            
            row_idx += 1
        
        # Mahal Ara Toplam
        mahal_end_row = row_idx - 1
        first_cell = xl_rowcol_to_cell(mahal_start_row, col_toplam)
        last_cell = xl_rowcol_to_cell(mahal_end_row, col_toplam)
        subtotal_formula = f"=SUM({first_cell}:{last_cell})"
        
        label_col = col_toplam - 1
        worksheet.write(row_idx, label_col, f"{mahal} TOPLAMI", fmt_mahal_toplam_label)
        worksheet.write_formula(row_idx, col_toplam, subtotal_formula, fmt_mahal_toplam_val)
        
        subtotal_cell_ref = xl_rowcol_to_cell(row_idx, col_toplam)
        mahal_ara_toplam_hucreleri.append(subtotal_cell_ref)
        
        row_idx += 1 

    # --- 6. GENEL TOPLAMLAR ---
    row_idx += 2
    label_col = col_toplam - 1
    val_col = col_toplam
    
    if mahal_ara_toplam_hucreleri:
        grand_sum_formula = "=" + "+".join(mahal_ara_toplam_hucreleri)
    else:
        grand_sum_formula = "=0"

    if genel_iskonto_tutar > 0:
        worksheet.write(row_idx, label_col, "ARA TOPLAM", fmt_total_label)
        worksheet.write_formula(row_idx, val_col, grand_sum_formula, fmt_total_val, ara_toplam)
        cell_ara_toplam = xl_rowcol_to_cell(row_idx, val_col)
        row_idx += 1
        
        worksheet.write(row_idx, label_col, "GENEL İSKONTO", fmt_total_label)
        worksheet.write(row_idx, val_col, -genel_iskonto_tutar, fmt_total_val)
        cell_genel_iskonto = xl_rowcol_to_cell(row_idx, val_col)
        row_idx += 1
        
        final_formula = f"={cell_ara_toplam}+{cell_genel_iskonto}"
        worksheet.write(row_idx, label_col, "GENEL TOPLAM", fmt_total_label)
        worksheet.write_formula(row_idx, val_col, final_formula, fmt_total_val, genel_toplam)
        
    else:
        worksheet.write(row_idx, label_col, "GENEL TOPLAM", fmt_total_label)
        worksheet.write_formula(row_idx, val_col, grand_sum_formula, fmt_total_val, genel_toplam)
    
    # --- 7. NOTLAR ---
    row_idx += 2
    fmt_info = workbook.add_format({'bold': True, 'align': 'right', 'font_color': '#555555'})
    worksheet.write(row_idx, val_col, f"NAKLİYE: {nakliye_durum}", fmt_info)
    worksheet.write(row_idx+1, val_col, f"MONTAJ: {montaj_durum}", fmt_info)
    
    row_idx += 3
    worksheet.merge_range(row_idx, 0, row_idx, 4, "NOTLAR:", fmt_baslik_proje)
    worksheet.merge_range(row_idx+1, 0, row_idx+5, 6, notlar_text, fmt_left)

    workbook.close()
    return output.getvalue()

# ==============================================================================
# 9. MODAL PENCERELER
# ==============================================================================
@st.dialog("Müşteri Ekle")
def ekleme_penceresi():
    st.write("Hızlı Müşteri Ekleme")
    with st.form("ekle_form"):
        f_adi = st.text_input("Firma Adı")
        y_kisi = st.text_input("Yetkili Kişi")
        ads = st.text_area("Adres")
        if st.form_submit_button("Kaydet", type="primary"):
            if f_adi:
                musteri_ekle(f_adi, y_kisi, ads)
                st.session_state.secili_firma_adi = f_adi 
                st.success("Kaydedildi!")
                st.rerun()

@st.dialog("Düzenle")
def musteri_duzenle_penceresi(id, f, y, a):
    with st.form("duzenle_form"):
        yeni_ad = st.text_input("Firma Adı", value=f)
        yeni_yetkili = st.text_input("Yetkili Kişi", value=y)
        yeni_adres = st.text_area("Adres", value=a)
        if st.form_submit_button("Güncelle", type="primary"):
            musteri_guncelle(id, yeni_ad, yeni_yetkili, yeni_adres)
            if st.session_state.secili_firma_adi == f:
                st.session_state.secili_firma_adi = yeni_ad
            st.success("Bilgiler güncellendi!")
            st.rerun()

@st.dialog("Sil")
def silme_onay_penceresi(id, f):
    st.warning(f"{f} silinecek.")
    c1, c2 = st.columns(2)
    if c1.button("Sil"):
        musteri_sil(id)
        if st.session_state.secili_firma_adi == f:
             st.session_state.secili_firma_adi = None 
        st.rerun() 
    if c2.button("İptal"):
        st.rerun()

@st.dialog("Teklif Sil")
def teklif_sil_onay_penceresi(id, p, r):
    st.error(f"🚨 **{p}** - **{r}** siliniyor.")
    
    # --- EKSİK OLAN KISIM BURASIYDI ---
    c1, c2 = st.columns(2)
    # ----------------------------------

    if c1.button("🗑️ Sil", type="primary"):
        teklif_sil(id)
        st.success("Silindi!")
        st.rerun()
        
    if c2.button("İptal"):
        st.rerun()

# --- GÖRSEL PENCERESİ (ÖNİZLEME + YENİDEN DÜZENLEME + SİLME) ---
@st.dialog("Görsel Düzenleyici", width="large")
def gorsel_penceresi(t_no):
    # 1. İlgili satırı bul
    df = st.session_state.tablo_verisi
    try:
        idx = df[df["NO"] == t_no].index[0]
    except:
        st.error("Ürün bulunamadı.")
        return

    st.markdown(f"### 📂 Ürün No: **{t_no}**")

    # 2. MEVCUT RESMİ HAFIZAYA AL (Yeniden Düzenleme Özelliği)
    # Eğer geçici hafızada (temp_img) resim yoksa ama veritabanında varsa, onu yükle.
    if st.session_state.temp_img is None:
        mevcut_resim = df.at[idx, "GÖRSEL"]
        if mevcut_resim and len(str(mevcut_resim)) > 50:
            try:
                st.session_state.temp_img = base64_to_image(mevcut_resim)
            except:
                pass # Hatalı data varsa geç

    # 3. Yeni Dosya Yükleyici
    up = st.file_uploader("Yeni Resim Yükle (JPG/PNG)", type=["jpg", "png", "jpeg"], key="img_uploader_modal")

    # Dosya seçilirse hafızayı güncelle
    if up:
        img = Image.open(up)
        img = jpeg_icin_hazirla(img)
        st.session_state.temp_img = img

    # --- EKRAN DÜZENİ (SOL: KIRPMA | SAĞ: ÖNİZLEME VE BUTONLAR) ---
    col_crop, col_preview = st.columns([2, 1], gap="large")

    cropped_img = None

    # SOL SÜTUN: KIRPMA ARACI
    with col_crop:
        if st.session_state.temp_img is not None:
            st.info("✂️ Sol tarafta mavi çerçeveyi ayarlayın.")
            # Kırpma Aracı
            cropped_img = st_cropper(
                st.session_state.temp_img, 
                realtime_update=True, # Canlı güncelleme açık
                box_color='#0000FF', 
                aspect_ratio=None,
                key="crop_widget"
            )
        else:
            st.warning("⚠️ Düzenlenecek görsel yok. Lütfen yukarıdan dosya seçin.")

    # SAĞ SÜTUN: ÖNİZLEME VE İŞLEMLER
    with col_preview:
        st.markdown("#### 👀 Önizleme")
        
        # Önizleme Kutusu
        if cropped_img:
            # Önizlemeyi biraz küçültüp gösterelim ki sığsın
            preview_show = cropped_img.copy()
            preview_show.thumbnail((300, 300)) 
            st.image(preview_show, caption="Eklenecek Görsel", use_container_width=True)
            st.success("Görünüm uygunsa kaydedin.")
        else:
            st.info("Kırpma alanı bekleniyor...")
            
        st.markdown("---")
        st.markdown("#### İşlemler")

        # 1. KAYDET BUTONU
        # Eğer resim varsa buton aktif olur
        if st.session_state.temp_img is not None:
            if st.button("💾 Kırp ve Kaydet", type="primary", use_container_width=True):
                if cropped_img:
                    buf = io.BytesIO()
                    # Kırpılan resmi kaydet
                    cropped_img.convert("RGB").save(buf, format="JPEG", quality=90)
                    b64 = base64.b64encode(buf.getvalue()).decode()
                    
                    # Veriyi tabloya işle
                    st.session_state.tablo_verisi.at[idx, "GÖRSEL"] = f"data:image/jpeg;base64,{b64}"
                    
                    # Temizlik
                    st.session_state.temp_img = None
                    st.session_state.aggrid_key += 1 
                    st.success("Kaydedildi!")
                    time.sleep(0.5)
                    st.rerun()

        # 2. SİL BUTONU
        if st.button("🗑️ Resmi Sil", type="secondary", use_container_width=True):
            st.session_state.tablo_verisi.at[idx, "GÖRSEL"] = ""
            st.session_state.temp_img = None
            st.session_state.aggrid_key += 1 
            st.warning("Resim kaldırıldı.")
            time.sleep(0.5)
            st.rerun()

        # 3. KAPAT BUTONU
        if st.button("❌ Vazgeç / Kapat", use_container_width=True):
            st.session_state.temp_img = None
            st.rerun()

@st.dialog("Detay Düzenle")
def detay_duzenle_penceresi(t_no):
    # 1. İlgili satırı bul
    df = st.session_state.tablo_verisi
    try:
        # t_no integer gelmeli, garantiye alalım
        idx = df[df["NO"] == int(t_no)].index[0]
    except:
        st.error("Satır bulunamadı.")
        return

    st.write(f"📝 Düzenlenen Satır No: **{t_no}**")
    
    # 2. Mevcut Açıklamayı Getir
    # Eğer hücre boşsa string olarak '' gelsin, nan hatası vermesin
    mevcut_aciklama = str(df.at[idx, "AÇIKLAMA"]) if pd.notna(df.at[idx, "AÇIKLAMA"]) else ""
    
    # 3. Düzenleme Alanı
    yeni_aciklama = st.text_area(
        "Ürün Açıklaması / Teknik Detay", 
        value=mevcut_aciklama, 
        height=300,
        placeholder="Ürün özelliklerini buraya girebilirsiniz..."
    )
    
    col1, col2 = st.columns([1, 1])
    
    with col1:
        if st.button("💾 Kaydet", type="primary", use_container_width=True):
            # Veriyi güncelle
            st.session_state.tablo_verisi.at[idx, "AÇIKLAMA"] = yeni_aciklama
            
            # State temizliği
            st.session_state.aktif_detay_id = None 
            st.session_state.aggrid_key += 1 # Tabloyu yenilemeye zorla
            
            st.success("Kaydedildi!")
            time.sleep(0.3)
            st.rerun()
            
    with col2:
        if st.button("❌ Vazgeç", use_container_width=True):
            st.session_state.aktif_detay_id = None
            st.rerun()

# ==============================================================================
# 10. ANA UYGULAMA AKIŞI
# ==============================================================================
tablolari_olustur()
sema_kontrol()

# --- SOL MENÜ ---
with st.sidebar:
    # --- LOGO ALANI ---
    # Eğer klasörde 'logo.png' varsa onu göster, yoksa yazıyı göster
    if os.path.exists("logo2.png"):
        st.image("logo2.png", use_container_width=True)
    else:
        st.title("TEKLİF PRO")
        
    st.markdown("---")
    
    if st.button("🏠 Ana Sayfa", use_container_width=True): 
        st.session_state.update(
            sayfa_secimi="🏠 Ana Sayfa", 
            aktif_teklif_data=None, 
            islem_turu="yeni", 
            tablo_verisi=None, 
            form_proje="", 
            form_no="", 
            form_rev="", 
            teklif_notlari=""
        )
        st.rerun()
    # SOL MENÜ KISMI
    if st.button("📝 Teklif Hazırla", use_container_width=True): 
        st.session_state.sayfa_secimi = "📝 Teklif Hazırla" # İsim tam olarak bu olmalı
        st.rerun()
    if st.button("🗂️ Geçmiş", use_container_width=True): 
        st.session_state.sayfa_secimi = "🗂️ Teklif Geçmişi"
        st.rerun()
    if st.button("👥 Müşteri Yönetimi (CRM)", use_container_width=True): 
        st.session_state.sayfa_secimi = "👥 Müşteri Yönetimi"
        st.rerun()
    if st.button("📜 Sözleşmeler (Resmi)", use_container_width=True):
        st.session_state.sayfa_secimi = "📜 Sözleşmeler"
        st.rerun()
    if st.button("🚛 Teslim Tutanağı", use_container_width=True):
        st.session_state.sayfa_secimi = "🚛 Teslim Tutanağı"
        st.rerun()
    if st.button("⚙️ Sistem", use_container_width=True):
        st.session_state.sayfa_secimi = "⚙️ Sistem"
        st.rerun()    
    
    st.markdown("---")
    if st.session_state.sayfa_secimi in ["📝 Teklif Hazırla", "🗂️ Teklif Geçmişi"]:
        m_df = musterileri_getir()
        if not m_df.empty:
            firms = m_df["firma_adi"].tolist()
            if st.session_state.secili_firma_adi in firms:
                idx = firms.index(st.session_state.secili_firma_adi)
            else:
                idx = None
            
            sel = st.selectbox("Firma Seç", firms, index=idx, key="sb_musteri")
            if sel:
                st.session_state.secili_firma_adi = sel
                secilen_firma_data = m_df[m_df["firma_adi"] == sel].iloc[0]
                
            # Taslak Listesi
            if 'secilen_firma_data' in locals() and secilen_firma_data is not None:
                t_df = taslaklari_getir(int(secilen_firma_data['id']))
                if not t_df.empty:
                    st.markdown("---")
                    st.caption("Taslaklar")
                    
                    taslak_secenekleri = {}
                    for index, row in t_df.iterrows():
                        try: tutar_val = float(row['toplam_tutar'])
                        except: tutar_val = 0.0
                        p_sembol = row['para_birimi'] if row['para_birimi'] else "TL"
                        p_adi_goster = row['proje_adi'] if row['proje_adi'] else "İsimsiz Proje"
                        etiket = f"{p_adi_goster} | {row['tarih']} | {tutar_val:,.0f} {p_sembol}"
                        taslak_secenekleri[etiket] = row

                    secilen_taslak_etiket = st.selectbox(
                        "Taslak Seçiniz:",
                        list(taslak_secenekleri.keys()),
                        key="sb_taslak_listesi",
                        label_visibility="collapsed",
                        placeholder="Taslak Seç..."
                    )
                    
                    if secilen_taslak_etiket:
                        secilen_taslak_data = taslak_secenekleri[secilen_taslak_etiket]
                        
                        col_yukle, col_sil = st.columns([4, 1])
                        with col_yukle:
                            if st.button("📂 Yükle", key="btn_taslak_yukle", use_container_width=True):
                                st.session_state.aktif_teklif_data = secilen_taslak_data
                                st.session_state.islem_turu = "taslak_duzenle"
                                st.session_state.tablo_verisi = None
                                st.session_state.aktif_taslak_id = secilen_taslak_data['id']
                                st.session_state.sayfa_secimi = "📝 Teklif Hazırla"
                                st.rerun()
                        
                        with col_sil:
                            if st.button("🗑️", key="btn_taslak_sil", help="Taslağı Sil", use_container_width=True):
                                teklif_sil_onay_penceresi(secilen_taslak_data['id'], secilen_taslak_data['proje_adi'], secilen_taslak_data['revizyon'])
        else:
            st.warning("Müşteri Yok")
        
        if st.button("➕ Ekle"):
            ekleme_penceresi()

    st.markdown("<div style='position:fixed; bottom:0; padding:10px; color:#95a5a6; font-size:10px;'>Saraks Furniture | Interior Fit-Out | Contracting</div>", unsafe_allow_html=True)

# ==============================================================================
# 11. SAYFA İÇERİKLERİ
# ==============================================================================

if st.session_state.sayfa_secimi == "🏠 Ana Sayfa":
    
    st.markdown(f"""
    <div class="dashboard-card" style="background: linear-gradient(135deg, #1E3A8A 0%, #2563EB 100%); color: white; border:none; margin-bottom:20px;">
        <h1 style="color: white; margin-bottom: 5px;">Hoş Geldiniz 👋</h1>
        <p style="opacity: 0.9; margin: 0;">Saraks Mobilya - Profesyonel Teklif Yönetim Sistemi</p>
    </div>
    """, unsafe_allow_html=True)
    
    # Kur Kontrolü
    if st.session_state.doviz_kurlari.get("USD", 0) == 0:
        yeni_kurlar = kurlari_getir()
        if yeni_kurlar.get("USD", 0) != 0:
            st.session_state.doviz_kurlari = yeni_kurlar
            st.rerun()
        else:
            st.warning("⚠️ Otomatik kur çekilemedi (Erişim engeli). Lütfen manuel giriniz.")
            c1, c2, c3 = st.columns(3)
            usd = c1.number_input("USD Kuru", value=36.0)
            eur = c2.number_input("EUR Kuru", value=39.0)
            gbp = c3.number_input("GBP Kuru", value=46.0)
            if st.button("Kurları Kaydet"):
                st.session_state.doviz_kurlari = {"TL": 1.0, "USD": usd, "EUR": eur, "GBP": gbp}
                st.rerun()
    
    k = st.session_state.doviz_kurlari
    if k.get("USD", 0) > 0:
        c1, c2, c3 = st.columns(3)
        c1.metric("🇺🇸 USD / TRY", f"₺{k['USD']:.2f}", delta_color="normal")
        c2.metric("🇪🇺 EUR / TRY", f"₺{k['EUR']:.2f}", delta_color="normal")
        c3.metric("🇬🇧 GBP / TRY", f"₺{k['GBP']:.2f}", delta_color="normal")
    
    st.markdown("---")
    
    conn = db_baglan()
    try:
        toplam_teklif = pd.read_sql_query("SELECT count(*) as sayi FROM teklifler", conn).iloc[0]['sayi']
        toplam_musteri = pd.read_sql_query("SELECT count(*) as sayi FROM musteriler", conn).iloc[0]['sayi']
        son_teklifler = pd.read_sql_query("SELECT proje_adi, firma_adi, tarih, toplam_tutar, para_birimi FROM teklifler WHERE durum='Yayında' ORDER BY id DESC LIMIT 5", conn)
    except:
        toplam_teklif = 0
        toplam_musteri = 0
        son_teklifler = pd.DataFrame()
    conn.close()
    
    col_main, col_side = st.columns([2, 1])
    
    with col_main:
        st.markdown("### 🚀 Hızlı Erişim")
        if st.button("📝 Yeni Teklif Oluştur", type="primary", use_container_width=True):
            st.session_state.sayfa_secimi = "📝 Teklif Hazırla"
            st.session_state.islem_turu = "yeni"
            st.session_state.aktif_teklif_data = None
            st.session_state.tablo_verisi = None
            st.session_state.form_proje=""
            st.session_state.form_no=""
            st.session_state.form_rev=""
            st.session_state.teklif_notlari=""
            st.session_state.form_proje_kodu="KOD"
            st.rerun()
            
        if st.button("👥 Müşteri Ekle", use_container_width=True):
            ekleme_penceresi()
        
    with col_side:
        st.markdown("### 📊 Özet")
        st.metric("Toplam Müşteri", toplam_musteri)
        
elif st.session_state.sayfa_secimi == "📝 Teklif Hazırla":
    
    mod_text = "Yeni Teklif Oluştur"
    if st.session_state.islem_turu == "duzenle":
        mod_text = "Mevcut Teklifi Düzenle"
        st.warning("⚠️ **Düzenleme Modu:** Değişiklikler mevcut kaydın üzerine yazılır.")
    elif st.session_state.islem_turu == "revize":
        mod_text = "Revizyon Oluştur"
        st.info("ℹ️ **Revizyon Modu:** Eski teklif korunur, yeni sürüm oluşturulur.")
    elif st.session_state.islem_turu == "taslak_duzenle":
        mod_text = "Taslak Düzenleniyor"
        st.info("💾 **Taslak Modu:** Çalışmalarınız otomatik kaydediliyor.")
    
    col_t, col_k = st.columns([2, 1])
    col_t.title(f"📝 {mod_text}")
    kurlar = st.session_state.doviz_kurlari
    if kurlar.get("USD", 0) > 0:
        col_k.success(f"Dolar: {kurlar.get('USD',0):.2f} | Euro: {kurlar.get('EUR',0):.2f}")

    if st.session_state.secili_firma_adi is not None:
        # Seçilen Firma Datası
        m_df_temp = musterileri_getir()
        secilen_firma_data = m_df_temp[m_df_temp["firma_adi"] == st.session_state.secili_firma_adi].iloc[0]

        # Form Değerleri Başlangıç
        def_proje = ""
        def_no = ""
        def_rev = ""
        def_tarih = date.today()
        def_para = "TL"
        def_kdv = 0

        # Tablo Verisi Hazırlığı
        if st.session_state.tablo_verisi is None:
            # SIFIR BAŞLANGIÇ İÇİN BOŞ DF
            init_df = pd.DataFrame({
                "SİL": [False], "MAHAL": ["Genel"], "NO": [1], "GÖRSEL_DURUM": [False], "DETAY_DURUM": [False], 
                "KOD": ["KOD-01"], "GÖRSEL": [""], 
                "ÜRÜN ADI": [""], 
                "AÇIKLAMA": [""], 
                "ÖLÇÜ": [""],
                "MİKTAR": [1], "BİRİM": ["Adet"], "BİRİM FİYAT": [0.0], 
                "İSKONTO": [0], "TOPLAM FİYAT": [0.0]
            })
            st.session_state.tablo_aktif_para_birimi = "TL"

            if st.session_state.aktif_teklif_data is not None:
                eski = st.session_state.aktif_teklif_data
                def_proje = eski["proje_adi"]
                def_no = eski["proje_no"]
                if "para_birimi" in eski and eski["para_birimi"]: 
                    def_para = eski["para_birimi"]
                    st.session_state.tablo_aktif_para_birimi = def_para
                if "kdv_orani" in eski and eski["kdv_orani"]: def_kdv = int(eski["kdv_orani"])
                if st.session_state.islem_turu == "revize": def_rev = sonraki_revizyon(eski["revizyon"]) 
                else: def_rev = eski["revizyon"] 
                try: def_tarih = datetime.strptime(eski["tarih"], '%Y-%m-%d').date()
                except: 
                    try: def_tarih = datetime.strptime(eski["tarih"], '%d-%m-%Y').date()
                    except: def_tarih = date.today()
                
                if "ozel_notlar" in eski and eski["ozel_notlar"]:
                    st.session_state.teklif_notlari = eski["ozel_notlar"]
                
                if "genel_iskonto" in eski and eski["genel_iskonto"]:
                    st.session_state.genel_iskonto = float(eski["genel_iskonto"])
                else:
                    st.session_state.genel_iskonto = 0.0

                if "nakliye_durum" in eski: st.session_state.nakliye_secimi = eski["nakliye_durum"]
                if "montaj_durum" in eski: st.session_state.montaj_secimi = eski["montaj_durum"]
                if "proje_kodu" in eski and eski["proje_kodu"]: st.session_state.form_proje_kodu = eski["proje_kodu"]

                if eski["urun_datasi"]:
                    init_df = pd.read_json(io.StringIO(eski["urun_datasi"]))
                    if "📸" in init_df.columns: init_df.rename(columns={"📸": "GÖRSEL_DURUM"}, inplace=True)
                    init_df.rename(columns=lambda x: x.replace(" ($)", "").replace(" (TL)", ""), inplace=True)
                    if "NO" not in init_df.columns: init_df.insert(0, "NO", range(1, len(init_df) + 1))
                    if "MAHAL" not in init_df.columns: init_df.insert(1, "MAHAL", "Genel")
                    if "İSKONTO" not in init_df.columns: init_df["İSKONTO"] = 0 
                    
                    for col in ["SİL", "MAHAL", "NO", "GÖRSEL_DURUM", "DETAY_DURUM", "KOD", "GÖRSEL", "ÜRÜN ADI", "AÇIKLAMA", "ÖLÇÜ", "MİKTAR", "BİRİM", "BİRİM FİYAT", "İSKONTO", "TOPLAM FİYAT"]:
                        if col not in init_df.columns:
                            if col == "SİL": init_df[col] = False
                            elif col == "MAHAL": init_df[col] = "Genel"
                            elif col == "GÖRSEL_DURUM": init_df[col] = False
                            elif col == "DETAY_DURUM": init_df[col] = False
                            elif col == "MİKTAR": init_df[col] = 1
                            elif col == "BİRİM": init_df[col] = "Adet"
                            elif col == "BİRİM FİYAT": init_df[col] = 0.0
                            elif col == "İSKONTO": init_df[col] = 0
                            elif col == "TOPLAM FİYAT": init_df[col] = 0.0
                            elif col == "NO": init_df[col] = range(1, len(init_df) + 1)
                            elif col == "AÇIKLAMA": init_df[col] = ""
                            else: init_df[col] = ""
            
            st.session_state.tablo_verisi = init_df
            st.session_state.form_proje = def_proje
            st.session_state.form_no = def_no
            st.session_state.form_rev = def_rev
            st.session_state.form_tarih = def_tarih
            st.session_state.form_para = def_para
            st.session_state.form_kdv = def_kdv

        # --- SEKMELİ YAPI ---
        tab_proje, tab_urun = st.tabs(["1. Proje Detayları", "2. Ürün Listesi & Hesaplama"])
        
        with tab_proje:
            with st.container(border=True):
                st.markdown("##### 🏗️ Proje Bilgileri")
                c1, c2, c3, c4, c5 = st.columns(5)
                p_adi = c1.text_input("Proje Adı", value=st.session_state.form_proje if st.session_state.form_proje else def_proje)
                
                # --- PROJE NO DEĞİŞİKLİĞİ VE OTOMATİK KOD GÜNCELLEME TETİKLEYİCİSİ ---
                proje_no_yeni = c2.text_input("Proje No (KOD Prefix)", value=st.session_state.form_no if st.session_state.form_no else def_no)
                
                if proje_no_yeni != st.session_state.form_no:
                    st.session_state.form_no = proje_no_yeni
                    # Proje no değişirse kodları hemen güncelle
                    if st.session_state.tablo_verisi is not None:
                        st.session_state.tablo_verisi = kodlari_yeniden_sirala(st.session_state.tablo_verisi, st.session_state.form_no)
                        st.session_state.aggrid_key += 1
                        st.rerun()
                
                # KOD PREFIX AYARI (ÖZEL KOD İÇİN)
                p_kodu = c3.text_input("Kod Öneki (Opsiyonel)", value=st.session_state.form_proje_kodu, help="Eğer boş bırakılırsa Proje No kullanılır.")
                
                p_rev = c4.text_input("Revizyon", value=st.session_state.form_rev if st.session_state.form_rev else def_rev)
                # TARIH FORMATI (DD/MM/YYYY)
                st.session_state.form_tarih = c5.date_input("Tarih", st.session_state.form_tarih, format="DD/MM/YYYY")
                
                st.session_state.form_proje = p_adi
                st.session_state.form_no = proje_no_yeni
                st.session_state.form_rev = p_rev
                st.session_state.form_proje_kodu = p_kodu

            with st.container(border=True):
                st.markdown("##### 💰 Para Birimi & KDV & Dil")
                c1, c2, c3 = st.columns(3)
                # DÖVİZ ÇEVİRİ BUTONU İÇİN ÖNCE SELECTBOX
                # DEĞİŞİKLİK: secilen_para -> pb
                pb = c1.selectbox("Para Birimi", ["TL", "USD", "EUR", "GBP"], index=["TL","USD","EUR","GBP"].index(st.session_state.tablo_aktif_para_birimi))
                
                # DEĞİŞİKLİK: secilen_kdv -> kdv
                kdv = c2.selectbox("KDV (%)", [0, 10, 20], index=[0,10,20].index(def_kdv))
                
                # DEĞİŞİKLİK: sec_dil -> dil
                dil = c3.selectbox("Teklif Dili / Language", ["TR", "EN"], index=0 if st.session_state.secili_dil == "TR" else 1)
                st.session_state.secili_dil = dil
                
                c4, c5 = st.columns(2)
                nakliye_opsiyonlari = ["HARİÇ", "DAHİL"]
                try: n_idx = nakliye_opsiyonlari.index(st.session_state.nakliye_secimi)
                except: n_idx = 0
                
                # DEĞİŞİKLİK: secilen_nakliye -> nak
                nak = c4.selectbox("Nakliye Durumu", nakliye_opsiyonlari, index=n_idx)
                st.session_state.nakliye_secimi = nak
                
                montaj_opsiyonlari = ["DAHİL", "HARİÇ"]
                try: m_idx = montaj_opsiyonlari.index(st.session_state.montaj_secimi)
                except: m_idx = 0
                
                # DEĞİŞİKLİK: secilen_montaj -> mon
                mon = c5.selectbox("Montaj Durumu", montaj_opsiyonlari, index=m_idx)
                st.session_state.montaj_secimi = mon

                # --- ÇEVİRİ MANTIĞI BURADA ---
                # Session State'i hemen güncelleme, önce kontrol et
                mevcut_para = st.session_state.tablo_aktif_para_birimi
                
                if pb != mevcut_para:
                    kaynak_kur = st.session_state.doviz_kurlari.get(mevcut_para, 1.0)
                    hedef_kur = st.session_state.doviz_kurlari.get(pb, 1.0)
                    
                    if hedef_kur == 0:
                         st.error("⚠️ Hata: Kur verisi 0. Lütfen internet bağlantısını kontrol edin veya ana sayfadan manuel girin.")
                    else:
                         carpan = kaynak_kur / hedef_kur
                         st.info(f"Tablo: **{mevcut_para}** ➔ Hedef: **{pb}**")
                         if st.button(f"🔄 Rakamları Dönüştür (x{carpan:.4f})", type="secondary"):
                            # DÜZELTME: Verileri önce sayıya çevir sonra çarp
                            df_convert = st.session_state.tablo_verisi.copy()
                            df_convert = temizle_ve_sayiya_cevir(df_convert, ["BİRİM FİYAT", "TOPLAM FİYAT", "MİKTAR", "İSKONTO"])
                            
                            df_convert["BİRİM FİYAT"] = df_convert["BİRİM FİYAT"] * carpan
                            df_convert["TOPLAM FİYAT"] = df_convert["MİKTAR"] * df_convert["BİRİM FİYAT"] * (1 - df_convert["İSKONTO"]/100)
                            
                            st.session_state.tablo_verisi = df_convert
                            st.session_state.tablo_aktif_para_birimi = pb
                            st.session_state.aggrid_key += 1
                            st.success(f"Dönüştürüldü!")
                            st.rerun()
                else:
                    # Eşitse güncelle (örneğin ilk açılışta veya iptal durumunda)
                    st.session_state.tablo_aktif_para_birimi = pb

        with tab_urun:
            # --- AgGrid KURULUMU (CHECKBOX FIX) ---
            
            # 1. Veriyi hazırla (Resim sütununu temizle - görsel olarak)
            df_grid = st.session_state.tablo_verisi.copy()
            
            # Görsel sütununda uzun base64 kodları yerine "VAR/YOK" gösterelim ki tablo donmasın
            df_grid["GÖRSEL_GRID"] = df_grid["GÖRSEL"].apply(lambda x: "📸 VAR" if x and len(str(x))>50 else "")
            
            # 2. Grid Yapılandırıcı (Excel Özellikleri)
            gb = GridOptionsBuilder.from_dataframe(df_grid)
            
            # GENEL AYARLAR
            gb.configure_default_column(editable=True, groupable=True, value=True, enableRowGroup=True, aggFunc='sum')
            
            # --- ÖNEMLİ: AUTO HEIGHT ---
            # SÜRÜKLE BIRAK YÖNETİMİ AKTİF (rowDragManaged=True)
            gb.configure_grid_options(
                domLayout='autoHeight', 
                rowDragManaged=True, 
                animateRows=True
            ) 
            
            # SÜTUN ÖZEL AYARLARI
            
            # --- OTOMATİK KOD SÜTUNU (KİLİTLİ) ---
            gb.configure_column("KOD", editable=False, cellStyle={'backgroundColor': '#f9f9f9', 'color': '#333'})
            
            gb.configure_column("NO", width=50, editable=False, rowDrag=True) # SÜRÜKLEME AKTİF
            
            # --- CHECKBOX FIX: KUTUCUĞU MAHAL SÜTUNUNA KOYUYORUZ ---
            gb.configure_column("MAHAL", width=120, checkboxSelection=True, headerCheckboxSelection=True)
            
            gb.configure_column("GÖRSEL", hide=True) # Gerçek base64 verisini gizle
            gb.configure_column("GÖRSEL_GRID", header_name="Görsel", width=90, editable=False, cellStyle={'textAlign': 'center'})
            
            gb.configure_column("ÜRÜN ADI", width=150)
            gb.configure_column("AÇIKLAMA", width=250, wrapText=True, autoHeight=True)
            gb.configure_column("ÖLÇÜ", width=100)
            gb.configure_column("MİKTAR", width=80, type=["numericColumn"])
            gb.configure_column("BİRİM FİYAT", width=110, type=["numericColumn"], valueFormatter="x.toLocaleString('tr-TR', {minimumFractionDigits: 2})")
            gb.configure_column("İSKONTO", width=80, type=["numericColumn"])
            
            # --- JS HESAPLAMA (Anlık Güncelleme) ---
            js_calc = JsCode("""
            function(params) {
                var miktar = params.data.MİKTAR || 0;
                var fiyat = params.data['BİRİM FİYAT'] || 0;
                var iskonto = params.data['İSKONTO'] || 0;
                return (miktar * fiyat) * (1 - iskonto / 100);
            }
            """)
            gb.configure_column("TOPLAM FİYAT", valueGetter=js_calc, width=120, editable=False, type=["numericColumn"], valueFormatter="x.toLocaleString('tr-TR', {minimumFractionDigits: 2})")
            
            # Button Columns (Gizli tutuyoruz, checkbox ile işlem yapacağız)
            gb.configure_column("GÖRSEL_DURUM", hide=True)
            gb.configure_column("DETAY_DURUM", hide=True)
            gb.configure_column("SİL", hide=True)

            # ÇOKLU SEÇİM (Checkbox)
            gb.configure_selection(selection_mode='multiple', use_checkbox=True)
            
            gridOptions = gb.build()

            # 3. TABLOYU ÇİZ
            st.info("ℹ️ Satırları sürükleyerek yerini değiştirebilir, sol baştaki kutucukla çoklu seçim yapabilirsiniz.")
            
            grid_response = AgGrid(
                df_grid, 
                gridOptions=gridOptions,
                update_mode=GridUpdateMode.MODEL_CHANGED | GridUpdateMode.SELECTION_CHANGED, # Hücre değişince tetikle
                data_return_mode=DataReturnMode.AS_INPUT, 
                fit_columns_on_grid_load=True,
                theme='streamlit', # 'streamlit', 'alpine', 'balham'
                height=None, # <-- ARTIK NONE YAPTIK Kİ AUTOHEIGHT ÇALIŞSIN
                key=f"grid_{st.session_state.aggrid_key}", # DINAMIK KEY
                allow_unsafe_jscode=True
            )

            # 4. GÜNCEL VERİYİ AL VE İŞLE
            updated_df = pd.DataFrame(grid_response['data'])
            selected_rows = grid_response['selected_rows']
            
            # Pandas DataFrame'e çevir
            if isinstance(updated_df, list):
                updated_df = pd.DataFrame(updated_df)
            
            # --- HESAPLAMA ---
            # Kullanıcı gridde sayıları değiştirdiyse toplamı güncelle
            # Önce sayısal tiplere zorla
            
            # HATA DÜZELTME: Verileri sayısal formata çevir (Temizlik Fonksiyonu ile)
            updated_df = temizle_ve_sayiya_cevir(updated_df, ["MİKTAR", "BİRİM FİYAT", "İSKONTO"])
            
            updated_df["TOPLAM FİYAT"] = updated_df["MİKTAR"] * updated_df["BİRİM FİYAT"] * (1 - updated_df["İSKONTO"]/100)
            
            # --- SIRA VE KOD GÜNCELLEME (SÜRÜKLEME SONRASI) ---
            if not updated_df.empty:
                # 1. Sıra Numaralarını (NO) 1'den başlayarak yeniden ver
                updated_df["NO"] = range(1, len(updated_df) + 1)
                
                # 2. Kodları bu yeni sıraya göre (Mutfak-01, Mutfak-02) yeniden oluştur
                # Prefix olarak ya özel girilen kodu ya da Proje No'yu kullan
                prefix_kullan = st.session_state.form_proje_kodu if st.session_state.form_proje_kodu else st.session_state.form_no
                updated_df = kodlari_yeniden_sirala(updated_df, prefix_kullan)
            
            # --- AUTO-SAVE MANTIĞI (V15 YENİLİĞİ) ---
            if not updated_df.equals(st.session_state.tablo_verisi) and secilen_firma_data is not None:
                ara_tmp = updated_df["TOPLAM FİYAT"].sum()
                genel_tmp = (ara_tmp * (1-st.session_state.genel_iskonto/100)) * (1+kdv/100)
                
                yeni_taslak_id = teklif_ekle_veya_guncelle(
                    st.session_state.aktif_taslak_id, 
                    int(secilen_firma_data['id']), 
                    secilen_firma_data['firma_adi'], 
                    st.session_state.form_proje, 
                    st.session_state.form_no, 
                    st.session_state.form_rev, 
                    str(st.session_state.form_tarih), 
                    genel_tmp, 
                    updated_df, 
                    pb, kdv, "Taslak", st.session_state.teklif_notlari, st.session_state.genel_iskonto, nak, mon, st.session_state.form_proje_kodu
                )
                st.session_state.aktif_taslak_id = yeni_taslak_id
                st.toast("Taslak Otomatik Kaydedildi", icon="💾")
            
            # Session State güncelle
            # NOT: Görsel sütunu AgGrid'de gizli olduğu için updated_df içinde var ama base64 verisi korunmuş olmalı.
            st.session_state.tablo_verisi = updated_df

            # --- BUTONLAR ---
            c_ekle, c_sil, c_img, c_det = st.columns([1, 1, 1.5, 1.5])
            
            with c_ekle:
                if st.button("➕ Satır Ekle", use_container_width=True):
                    # Benzersiz NO üretmek için max NO + 1
                    max_no = st.session_state.tablo_verisi["NO"].max()
                    yeni_no = max_no + 1 if not pd.isna(max_no) else 1
                    
                    yeni_satir = pd.DataFrame([{
                        "SİL":False, "MAHAL":"Genel", "NO":yeni_no, "GÖRSEL_DURUM":False, "DETAY_DURUM":False,
                        "KOD":"", "GÖRSEL":"", "GÖRSEL_GRID":"", "ÜRÜN ADI":"", "AÇIKLAMA":"", "ÖLÇÜ":"", 
                        "MİKTAR":1, "BİRİM":"Adet", "BİRİM FİYAT":0.0, "İSKONTO":0, "TOPLAM FİYAT":0.0
                    }])
                    st.session_state.tablo_verisi = pd.concat([st.session_state.tablo_verisi, yeni_satir], ignore_index=True)
                    
                    # Ekleme sonrası kodları yeniden düzenle
                    prefix_kullan = st.session_state.form_proje_kodu if st.session_state.form_proje_kodu else st.session_state.form_no
                    st.session_state.tablo_verisi = kodlari_yeniden_sirala(st.session_state.tablo_verisi, prefix_kullan)
                    
                    # KEY GÜNCELLE (Yenileme için)
                    st.session_state.aggrid_key += 1
                    st.rerun()

            with c_sil:
                if st.button("🗑️ Seçilileri Sil", type="primary", use_container_width=True):
                    if selected_rows is not None and len(selected_rows) > 0:
                        # Seçili satırların NO'larını al
                        if isinstance(selected_rows, pd.DataFrame):
                            # DataFrame ise
                             selected_list = selected_rows.to_dict('records')
                             nolar = selected_rows['NO'].tolist()
                        else:
                            # Liste ise
                             selected_list = selected_rows
                             nolar = [r['NO'] for r in selected_rows]
                        
                        # Bu NO'lara sahip OLMAYANLARI tut
                        st.session_state.tablo_verisi = st.session_state.tablo_verisi[~st.session_state.tablo_verisi['NO'].isin(nolar)].reset_index(drop=True)
                        
                        # Numaraları Yenile (Opsiyonel: Silince 1,2,3 diye tekrar dizilsin mi? Evet)
                        st.session_state.tablo_verisi["NO"] = range(1, len(st.session_state.tablo_verisi) + 1)
                        
                        # Silme sonrası kodları yeniden düzenle
                        prefix_kullan = st.session_state.form_proje_kodu if st.session_state.form_proje_kodu else st.session_state.form_no
                        st.session_state.tablo_verisi = kodlari_yeniden_sirala(st.session_state.tablo_verisi, prefix_kullan)

                        # KEY GÜNCELLE
                        st.session_state.aggrid_key += 1
                        st.success(f"{len(nolar)} satır silindi.")
                        st.rerun()
                    else:
                        st.warning("Listeden (Checkbox) seçim yapınız.")

            with c_img:
                if st.button("🖼️ Seçilinin Resmini Düzenle", use_container_width=True):
                    if selected_rows is not None and len(selected_rows) > 0:
                        # Seçilen ilk satırın NO bilgisini al
                        secilen_row = selected_rows.iloc[0] if isinstance(selected_rows, pd.DataFrame) else selected_rows[0]
                        secilen_no = secilen_row['NO']
                        
                        # HATA DÜZELTME: Hafızadaki eski resmi temizle
                        st.session_state.temp_img = None
                        
                        # Bu NO ile modal aç
                        gorsel_penceresi(secilen_no)
                    else:
                        st.warning("Lütfen bir satır seçin.")

            with c_det:
                if st.button("📝 Seçilinin Detayını Gir", use_container_width=True):
                    if selected_rows is not None and len(selected_rows) > 0:
                        secilen_row = selected_rows.iloc[0] if isinstance(selected_rows, pd.DataFrame) else selected_rows[0]
                        secilen_no = secilen_row['NO']
                        detay_duzenle_penceresi(secilen_no)
                    else:
                        st.warning("Lütfen bir satır seçin.")
                        
            # --- DIALOG ÇAKIŞMA KONTROLÜ (Eğer aktif detay varsa aç) ---
            if st.session_state.aktif_detay_id is not None:
                 detay_duzenle_penceresi(st.session_state.aktif_detay_id)

            # --- HESAPLAMALAR VE ALT PANEL YERLEŞİMİ ---
            edited = st.session_state.tablo_verisi # AgGrid ile güncellenen veri
            ara_top = edited["TOPLAM FİYAT"].sum()
            st.divider()

            c_notlar, c_ozet = st.columns([1.5, 1]) 

            with c_notlar:
                st.markdown("### 📝 Notlar")
                # Notlar alanını buraya taşıdık.
                notlar = st.text_area(
                    "Teklif Notları", 
                    value=st.session_state.teklif_notlari, 
                    height=300,
                    placeholder="Ödeme şartları, teslimat süresi vb. notları buraya girebilirsiniz..."
                )
                if notlar != st.session_state.teklif_notlari:
                    st.session_state.teklif_notlari = notlar

            with c_ozet:
                st.markdown("#### 🏘️ Bölüm Bazlı Özet")
                if not edited.empty:
                    mahal_gruplari = edited.groupby("MAHAL")["TOPLAM FİYAT"].sum().reset_index()
                    for _, satir in mahal_gruplari.iterrows():
                        st.metric(
                            label=f"📍 {satir['MAHAL']}", 
                            value=format_para(satir['TOPLAM FİYAT'], pb, 2)
                        )
                
                st.divider()

                # 1. Gerekli değişkenlerin ve katsayıların hazırlanması
                kdv_katsayi = 1 + (kdv / 100.0) # Örn: %20 KDV için 1.20

                # Session State başlatma (Eğer henüz yoksa)
                if "genel_iskonto" not in st.session_state:
                    st.session_state.genel_iskonto = 0.0

                # --- CALLBACK FONKSİYONLARI (İşin beyni burası) ---

                def on_iskonto_change():
                    """İskonto yüzdesi değişince çalışır, toplamı günceller."""
                    yeni_oran = st.session_state.w_iskonto
                    st.session_state.genel_iskonto = yeni_oran
                    
                    # Yeni Toplamı Hesapla: AraToplam * (1 - İskonto) * KDV
                    yeni_toplam = ara_top * (1 - yeni_oran / 100.0) * kdv_katsayi
                    st.session_state.w_toplam = yeni_toplam

                def on_toplam_change():
                    """Genel Toplam elle değiştirilince çalışır, iskonto oranını bulur."""
                    yeni_girilen_toplam = st.session_state.w_toplam
                    
                    if ara_top > 0:
                        # Formül: İskonto% = 100 * (1 - (HedefToplam / (AraToplam * KDV_Katsayisi)))
                        hedef_net = yeni_girilen_toplam / kdv_katsayi
                        hesaplanan_oran = 100 * (1 - (hedef_net / ara_top))
                        
                        # Sınır Kontrolü (Negatif iskonto veya %100 üzeri olmasın)
                        if hesaplanan_oran < 0: hesaplanan_oran = 0.0
                        if hesaplanan_oran > 100: hesaplanan_oran = 100.0
                        
                        st.session_state.genel_iskonto = hesaplanan_oran
                        st.session_state.w_iskonto = hesaplanan_oran
                    else:
                        st.session_state.w_iskonto = 0.0

                # --- ARAYÜZ (UI) KISMI ---

                # Önce güncel değerleri state'lere yükleyelim (Senkronizasyon)
                # Eğer kullanıcı henüz bir kutuya dokunmadıysa, mevcut veriyi yansıtalım.
                if "w_iskonto" not in st.session_state:
                    st.session_state.w_iskonto = float(st.session_state.genel_iskonto)

                if "w_toplam" not in st.session_state:
                    # Başlangıç toplamını hesapla
                    mevcut_toplam = ara_top * (1 - st.session_state.genel_iskonto / 100.0) * kdv_katsayi
                    st.session_state.w_toplam = mevcut_toplam
                else:
                    # Ara toplam değişmiş olabilir (ürün ekleyince), bu durumda iskonto oranını sabit tutup toplamı güncelleyelim
                    # Ancak manuel giriş anında çakışma olmaması için burayı kontrollü yapıyoruz.
                    # Basit yöntem: Her döngüde iskonto oranına sadık kalıp toplamı yeniden hesaplamak en güvenlisidir.
                    guncel_beklenen_toplam = ara_top * (1 - st.session_state.genel_iskonto / 100.0) * kdv_katsayi
                    # Eğer fark çok küçükse (kuruş farkı), kullanıcının girdiği değeri koru, yoksa (ürün eklendiyse) güncelle
                    if abs(st.session_state.w_toplam - guncel_beklenen_toplam) > 1.0: 
                        st.session_state.w_toplam = guncel_beklenen_toplam
                        st.session_state.w_iskonto = st.session_state.genel_iskonto

                c_ara, c_bos = st.columns([1, 1])
                with c_ara:
                    st.metric("GENEL ARA TOPLAM", format_para(ara_top, pb, 2))

                # İki kutuyu yan yana koyuyoruz
                col_isk, col_top = st.columns([1, 1])

                with col_isk:
                    st.number_input(
                        "Genel İskonto (%)",
                        min_value=0.0,
                        max_value=100.0,
                        step=1.0,
                        format="%.2f",
                        key="w_iskonto",           # State anahtarı
                        on_change=on_iskonto_change # Değişince çalışacak fonksiyon
                    )

                with col_top:
                    st.number_input(
                        f"GENEL TOPLAM ({pb})",
                        min_value=0.0,
                        step=100.0,               # 100 birim artırıp azaltma
                        format="%.2f",
                        key="w_toplam",           # State anahtarı
                        on_change=on_toplam_change # Değişince çalışacak fonksiyon
                    )

                # --- ALT BİLGİ HESAPLAMALARI (PDF ve Excel için değişkenleri güncelle) ---
                # Bu değişkenler (isk_tutar, genel_top) aşağıdaki "KAYDET" butonları için gereklidir.
                isk_oran = st.session_state.genel_iskonto
                isk_tutar = ara_top * (isk_oran / 100.0)
                net = ara_top - isk_tutar
                kdv_tutar = net * (kdv / 100.0)
                genel_top = st.session_state.w_toplam # Doğrudan kutudaki değeri alıyoruz

                # Detay Gösterimi (İsteğe bağlı, bilgi amaçlı)
                if isk_oran > 0:
                    st.caption(f"Detay: {format_para(ara_top, pb)} - {format_para(isk_tutar, pb)} (İnd.) + {format_para(kdv_tutar, pb)} (KDV) = {format_para(genel_top, pb)}")

                st.markdown("---")
            c_save, c_pdf, c_excel = st.columns([1, 1, 1])
            
            # DOSYA İSMİ (PROJE NO)
            dosya_ismi = str(st.session_state.form_no).strip() if st.session_state.form_no else "teklif"
            # Geçersiz karakterleri temizle
            dosya_ismi = "".join([c for c in dosya_ismi if c.isalnum() or c in (' ', '-', '_')]).rstrip()
            
            with c_save:
                if st.button("💾 YAYINLA (KAYDET)", type="primary", use_container_width=True):
                    # 1. Ana Teklifi Kaydet (Mevcut Fonksiyon)
                    tgt = st.session_state.aktif_teklif_data['id'] if st.session_state.islem_turu=="duzenle" else (st.session_state.aktif_taslak_id if st.session_state.islem_turu=="taslak_duzenle" else None)
                    
                    yeni_teklif_id = teklif_ekle_veya_guncelle(
                        tgt, 
                        int(secilen_firma_data['id']), 
                        secilen_firma_data['firma_adi'], 
                        st.session_state.form_proje, 
                        st.session_state.form_no, 
                        st.session_state.form_rev, 
                        str(st.session_state.form_tarih), 
                        genel_top, 
                        st.session_state.tablo_verisi, 
                        pb, 
                        kdv, 
                        "Yayında", 
                        st.session_state.teklif_notlari, 
                        isk_oran, 
                        nak, 
                        mon,
                        st.session_state.form_proje_kodu
                    )

                    # -----------------------------------------------------------
                    # 2. EKSİK OLAN PARÇA: ÜRÜNLERİ DETAYLI TABLOYA İŞLEME
                    # -----------------------------------------------------------
                    try:
                        conn = db_baglan()
                        c = conn.cursor()
                        
                        # Tablo yoksa oluştur (Garanti olsun)
                        c.execute('''CREATE TABLE IF NOT EXISTS teklif_satirlari
                                     (id INTEGER PRIMARY KEY AUTOINCREMENT,
                                      teklif_id INTEGER,
                                      urun_kodu TEXT,
                                      urun_adi TEXT,
                                      miktar TEXT,
                                      birim TEXT,
                                      birim_fiyat REAL,
                                      toplam_fiyat REAL)''')
                        
                        # Önce bu teklife ait eski satırları temizle (Güncelleme yapıyorsak çakışmasın)
                        c.execute("DELETE FROM teklif_satirlari WHERE teklif_id = ?", (yeni_teklif_id,))
                        
                        # Tablodaki dolu satırları al
                        dolu_satirlar = edited[edited["ÜRÜN ADI"].str.strip() != ""]
                        
                        kayit_sayisi = 0
                        for i, row in dolu_satirlar.iterrows():
                            # Güvenli veri dönüşümü
                            k = str(row.get('KOD', ''))
                            u = str(row.get('ÜRÜN ADI', ''))
                            m = str(row.get('MİKTAR', '1'))
                            b = str(row.get('BİRİM', 'Adet'))
                            try: bf = float(row.get('BİRİM FİYAT', 0))
                            except: bf = 0.0
                            try: tf = float(row.get('TOPLAM FİYAT', 0))
                            except: tf = 0.0

                            c.execute("INSERT INTO teklif_satirlari (teklif_id, urun_kodu, urun_adi, miktar, birim, birim_fiyat, toplam_fiyat) VALUES (?, ?, ?, ?, ?, ?, ?)",
                                      (yeni_teklif_id, k, u, m, b, bf, tf))
                            kayit_sayisi += 1
                        
                        conn.commit()
                        conn.close()
                        
                        st.success(f"✅ Başarıyla Yayınlandı! ({kayit_sayisi} kalem ürün sisteme işlendi)")
                        
                        # Teslim Tutanağı sayfası için bir ipucu verelim
                        st.caption("ℹ️ Bu teklifi artık 'Teslim Tutanağı' sayfasında seçebilirsiniz.")
                        
                    except Exception as e:
                        st.error(f"Detay Kayıt Hatası: {e}")
                    # -----------------------------------------------------------
                    
            with c_pdf:
                if st.button("📄 PDF OLUŞTUR", type="secondary", use_container_width=True):
                    with st.spinner('PDF hazırlanıyor...'):
                        pdata = create_pdf(
                            secilen_firma_data, 
                            {'adi':st.session_state.form_proje, 'no':st.session_state.form_no, 'rev':st.session_state.form_rev, 'tarih':st.session_state.form_tarih}, 
                            st.session_state.tablo_verisi, 
                            ara_top, ara_top*(isk_oran/100), genel_top, pb, kdv, 
                            st.session_state.teklif_notlari, nak, mon, dil
                        )
                        b64 = base64.b64encode(pdata).decode()
                        
                    st.toast('PDF Hazırlandı!', icon='📄')
                    
                    # 1. İndirme Linki
                    st.markdown(f'<a href="data:application/pdf;base64,{b64}" download="{dosya_ismi}.pdf" style="display:block;width:100%;padding:10px;text-align:center;background:#E74C3C;color:white;border-radius:8px;text-decoration:none;font-weight:bold;margin-bottom:10px;">⬇️ PDF İNDİR</a>', unsafe_allow_html=True)
                    
                    # 2. PDF ÖNİZLEME (YENİ ÖZELLİK 🚀)
                    with st.expander("👀 PDF Önizleme", expanded=True):
                        pdf_display = f'<iframe src="data:application/pdf;base64,{b64}" width="100%" height="600" type="application/pdf"></iframe>'
                        st.markdown(pdf_display, unsafe_allow_html=True)
            
            with c_excel:
                if st.button("📗 EXCEL İNDİR", type="secondary", use_container_width=True):
                    edata = create_excel(
                        secilen_firma_data, 
                        {'adi':st.session_state.form_proje, 'no':st.session_state.form_no, 'tarih':st.session_state.form_tarih, 'rev': st.session_state.form_rev}, 
                        st.session_state.tablo_verisi, 
                        ara_top, 
                        ara_top*(isk_oran/100), 
                        genel_top, 
                        pb, 
                        st.session_state.teklif_notlari,
                        nak,  # Yeni Eklendi
                        mon   # Yeni Eklendi
                    )
                    b64_xl = base64.b64encode(edata).decode()
                    st.markdown(f'<a href="data:application/vnd.openxmlformats-officedocument.spreadsheetml.sheet;base64,{b64_xl}" download="{dosya_ismi}.xlsx" style="display:block;width:100%;padding:10px;text-align:center;background:#27AE60;color:white;border-radius:5px;text-decoration:none;">⬇️ EXCEL İNDİR</a>', unsafe_allow_html=True)

    else:
        st.warning("Lütfen sol menüden müşteri seçin.")


# --- TEKLİF GEÇMİŞİ SAYFASI (EKSİKSE BURAYA YAPIŞTIRIN) ---

elif st.session_state.sayfa_secimi == "🗂️ Teklif Geçmişi":
    st.title("Teklif Geçmişi")
    
    # --- 1. SEÇİLİ MÜŞTERİYE GÖRE FİLTRELEME ---
    if secilen_firma_data is not None:
        target_id = int(secilen_firma_data['id'])
        
        # Veriyi Çek
        conn = db_baglan()
        try:
            # SQL sorgusunu garantiye alalım
            df = pd.read_sql("SELECT * FROM teklifler WHERE musteri_id = ? ORDER BY id DESC", conn, params=(target_id,))
        except Exception as e:
            st.error(f"Veri çekme hatası: {e}")
            df = pd.DataFrame()
        finally:
            conn.close()
        
        if not df.empty:
            # --- NORMAL LİSTELEME KODU ---
            for p in df['proje_no'].unique():
                revs = df[df['proje_no']==p].sort_values("revizyon", ascending=False)
                if revs.empty: continue
                lat = revs.iloc[0]
                tutar_str = format_para(lat['toplam_tutar'], lat['para_birimi'])
                
                with st.expander(f"📂 {lat['proje_adi']} ({p}) - Son: {lat['revizyon']} | {tutar_str}"):
                     for idx, row in revs.iterrows():
                        # Renk Ayarları
                        durum = row.get('durum', 'Yayında')
                        colors = {
                            "Taslak": ("#F1F5F9", "#475569", "📝"),
                            "Yayında": ("#DCFCE7", "#166534", "✅"),
                            "Onaylandı": ("#166534", "#FFFFFF", "🏆"),
                            "Reddedildi": ("#FEF2F2", "#DC2626", "❌"),
                            "Beklemede": ("#FEF9C3", "#854D0E", "⚠️")
                        }
                        bg, txt, icon = colors.get(durum, ("#F1F5F9", "#000000", "❓"))
                        
                        # Satır Yapısı (Container ile)
                        with st.container(border=True):
                            c1, c2, c3 = st.columns([3, 1.5, 1.5])
                            c1.markdown(f"**Rev: {row['revizyon']}** | 📅 {row['tarih']}")
                            c2.markdown(f'<span style="background:{bg}; color:{txt}; padding:4px 8px; border-radius:8px; font-size:12px;">{icon} {durum}</span>', unsafe_allow_html=True)
                            c3.markdown(f"**{format_para(row['toplam_tutar'], row['para_birimi'])}**")
                            
                            # Butonlar
                            b1, b2, b3, b4 = st.columns([1,1,1,0.5])
                            if b1.button("✏️ Düzenle", key=f"e{row['id']}", help="Düzenle"):
                                st.session_state.update(aktif_teklif_data=row, islem_turu="duzenle", tablo_verisi=None, sayfa_secimi="📝 Teklif Hazırla"); st.rerun()
                            if b2.button("📈 Revize Et", key=f"r{row['id']}", help="Revize"):
                                st.session_state.update(aktif_teklif_data=row, islem_turu="revize", tablo_verisi=None, sayfa_secimi="📝 Teklif Hazırla"); st.rerun()
                            if b3.button("📋 Kopyala", key=f"c{row['id']}", help="Kopyala"):
                                y = row.copy(); y['id']=None; y['revizyon']="00"; y['proje_adi']+=" (Kopya)"; y['tarih']=str(date.today())
                                st.session_state.update(aktif_teklif_data=y, islem_turu="yeni", tablo_verisi=None, sayfa_secimi="📝 Teklif Hazırla"); st.rerun()
                            if b4.button("🗑️ Sil", key=f"d{row['id']}"):
                                teklif_sil_onay_penceresi(row['id'], row['proje_adi'], row['revizyon'])
        else:
            st.warning(f"⚠️ '{secilen_firma_data['firma_adi']}' (ID: {target_id}) müşterisine ait kayıtlı teklif bulunamadı.")
            
    else:
        st.info("Lütfen sol menüden müşteri seçin.")

    st.markdown("---")
    
    # --- 2. DEBUG ALANI (Sorunu Çözecek Kısım) ---
    with st.expander("🔍 SİSTEM KONTROLÜ (Tüm Kayıtları Göster)", expanded=False):
        st.error("Eğer yukarıda tekliflerinizi göremiyorsanız buraya bakın.")
        
        conn = db_baglan()
        tum_kayitlar = pd.read_sql("SELECT * FROM teklifler", conn)
        conn.close()
        
        if not tum_kayitlar.empty:
            st.write(f"Veritabanında toplam **{len(tum_kayitlar)}** adet teklif var (Müşteri ayrımı yapmaksızın):")
            st.dataframe(tum_kayitlar)
            st.info("Eğer teklifiniz burada var ama yukarıda yoksa 'musteri_id' sütununu kontrol edin.")
        else:
            st.error("🚨 Veritabanı TAMAMEN BOŞ! Kaydetme işlemi başarısız oluyor.")
            st.write("Lütfen 'Teklif Hazırla' sayfasına gidip 'KAYDET' butonuna bastığınızdan emin olun.")

# --- MÜŞTERİ YÖNETİMİ (HATASIZ PRO UI 💎) ---
# --- MÜŞTERİ YÖNETİMİ (OTOMATİK KAYIT ÖZELLİĞİ EKLENDİ ✅) ---
# --- MÜŞTERİ YÖNETİMİ (TAM SÜRÜM - EKSİKSİZ KOD ✅) ---
elif st.session_state.sayfa_secimi == "👥 Müşteri Yönetimi":
    st.title("Müşteri İlişkileri Yönetimi (CRM)")
    
    # Verileri Çek
    df_musteriler = musterileri_getir() #
    
    # 1. SİSTEM GENELİ ÖZET
    conn = db_baglan()
    try:
        df_tum_teklifler = pd.read_sql("SELECT * FROM teklifler", conn)
        genel_projeler = df_tum_teklifler.drop_duplicates(subset=['proje_no'], keep='first')
        toplam_aktif = len(genel_projeler[genel_projeler['durum'].isin(['Onaylandı'])])
        toplam_musteri = len(df_musteriler)
    except:
        toplam_aktif = 0
        toplam_musteri = 0
    finally:
        conn.close()

    col_g1, col_g2 = st.columns(2)
    col_g1.metric("Toplam Kayıtlı Müşteri", toplam_musteri)
    col_g2.metric("Sistemdeki Aktif Proje Sayısı", toplam_aktif)
    
    st.divider()

    # 2. İŞLEM YAPILACAK MÜŞTERİ SEÇİMİ (ÜSTTE ✅)
    st.subheader("🎯 Müşteri Analizi ve Teklif Yönetimi")
    if not df_musteriler.empty:
        secilen_musteri = st.selectbox(
            "İşlem Yapılacak Müşteriyi Seçin:", 
            df_musteriler['firma_adi'].tolist(), 
            index=None,
            placeholder="Hızlıca bir müşteri seçin ve detaylarını görün...",
            label_visibility="collapsed"
        )
    else:
        secilen_musteri = None

    # --- SEÇİLEN MÜŞTERİ DETAY ALANI ---
    if secilen_musteri:
        musteri_row = df_musteriler[df_musteriler["firma_adi"] == secilen_musteri].iloc[0]
        m_id = int(musteri_row['id'])
        
        st.markdown(f"### 🏢 {musteri_row['firma_adi']} - Cari Kart")
        
        # Müşteri Bilgi Kartı
        with st.container(border=True):
            c_bilgi, c_aksiyon = st.columns([3, 1])
            with c_bilgi:
                # - Müşteri bilgileri gösterimi
                st.write(f"👤 **{musteri_row['yetkili_kisi']}** | 📞 {musteri_row.get('telefon', 'Belirtilmedi')}")
                st.caption(f"📍 {musteri_row['adres']}")
            with c_aksiyon:
                ca1, ca2 = st.columns(2)
                if ca1.button("✏️", help="Düzenle", use_container_width=True, key="m_edit"):
                    musteri_duzenle_penceresi(m_id, musteri_row['firma_adi'], musteri_row['yetkili_kisi'], musteri_row['adres'])
                if ca2.button("🗑️", help="Sil", type="primary", use_container_width=True, key="m_del"):
                    silme_onay_penceresi(m_id, musteri_row['firma_adi'])

        # Müşteriye Ait Teklifleri Çek
        conn = db_baglan()
        m_teklifler = pd.read_sql("SELECT * FROM teklifler WHERE musteri_id = ? ORDER BY id DESC", conn, params=(m_id,))
        conn.close()
        
        if not m_teklifler.empty:
            # YIL FİLTRESİ
            m_teklifler['yil_temp'] = pd.to_datetime(m_teklifler['tarih']).dt.year
            mevcut_yillar = sorted(m_teklifler['yil_temp'].unique().tolist(), reverse=True)
            yillar_listesi = ["Tüm Zamanlar"] + [str(y) for y in mevcut_yillar]
            
            c_f1, _ = st.columns([1, 4])
            secilen_yil = c_f1.selectbox("📅 Yıl Seçin", yillar_listesi, index=0)
            
            df_final = m_teklifler.copy()
            if secilen_yil != "Tüm Zamanlar":
                df_final = df_final[df_final['yil_temp'] == int(secilen_yil)]

            # YARDIMCI FONKSİYONLAR (Burada Tanımlanmalı)
            def get_html_val(df_in):
                if df_in.empty: return '<div style="color:#94A3B8;">0 ₺</div>'
                gruplu = df_in.groupby("para_birimi")["toplam_tutar"].sum()
                symbols = {"TL": "₺", "USD": "$", "EUR": "€"}
                out = ""
                for pb, tutar in gruplu.items():
                    s = symbols.get(pb, pb)
                    val = f"{tutar:,.0f} {s}".replace(",", "X").replace(".", ",").replace("X", ".")
                    out += f"<div>{val}</div>"
                return out

            def make_card(label, val_html, sub, color):
                return f"""<div style="background:white; border:1px solid #E2E8F0; border-top:4px solid {color}; border-radius:10px; padding:15px; height:100%;">
                    <div style="color:{color}; font-size:11px; font-weight:bold; text-transform:uppercase;">{label}</div>
                    <div style="font-size:16px; font-weight:700; margin-top:5px; color:#0F172A;">{val_html}</div>
                    <div style="color:#94A3B8; font-size:10px; margin-top:8px; border-top:1px solid #F1F5F9; padding-top:4px;">{sub}</div>
                </div>"""

            # HESAPLAMALAR
            en_guncel = df_final.drop_duplicates(subset=['proje_no'], keep='first')
            onay = en_guncel[en_guncel['durum'] == 'Onaylandı']
            bekle = en_guncel[en_guncel['durum'].isin(['Yayında', 'Revize', 'Beklemede'])]
            red = en_guncel[en_guncel['durum'] == 'Reddedildi']

            # 5'Lİ KPI KARTLARI
            st.write(f"##### 📊 {secilen_yil} Finansal Özeti")
            k1, k2, k3, k4, k5 = st.columns(5)
            with k1: st.markdown(make_card("Toplam Proje", f"{len(en_guncel)}", "Tekil Adet", "#64748B"), unsafe_allow_html=True)
            with k2: st.markdown(make_card("Onaylanan (Ciro)", get_html_val(onay), f"{len(onay)} Proje", "#166534"), unsafe_allow_html=True)
            with k3: st.markdown(make_card("Bekleyen", get_html_val(bekle), f"{len(bekle)} Proje", "#D97706"), unsafe_allow_html=True)
            with k4: st.markdown(make_card("Reddedilen (Kayıp)", get_html_val(red), f"{len(red)} Proje", "#DC2626"), unsafe_allow_html=True)
            with k5:
                basari = (len(onay)/len(en_guncel)*100) if len(en_guncel)>0 else 0
                st.markdown(make_card("Başarı Oranı", f"%{basari:.1f}", f"{len(red)} Kayıp", "#1E40AF"), unsafe_allow_html=True)

            st.divider()

            # PROJE LİSTESİ
            tab_liste, tab_grafik = st.tabs(["📂 Proje Detayları", "📈 Yıllık Performans"])
            
            with tab_liste:
                proje_nolar = df_final['proje_no'].unique()
                for p_no in proje_nolar:
                    revs = df_final[df_final['proje_no'] == p_no]
                    son = revs.iloc[0]
                    with st.expander(f"📂 {son['proje_adi']} ({p_no}) | {format_para(son['toplam_tutar'], son['para_birimi'])}"):
                        for idx, row in revs.iterrows():
                            c_rev, c_tutar, c_stat, c_git = st.columns([1, 2, 2, 1])
                            c_rev.write(f"**Rev: {row['revizyon']}**")
                            c_tutar.write(format_para(row['toplam_tutar'], row['para_birimi']))
                            
                            # Durum Güncelleme
                            durum_listesi = ["Beklemede", "Yayında", "Onaylandı", "Reddedildi"]
                            default_idx = durum_listesi.index(row['durum']) if row['durum'] in durum_listesi else 0
                            yeni_durum = c_stat.selectbox("Durum", durum_listesi, index=default_idx, key=f"crm_stat_{row['id']}", label_visibility="collapsed")
                            
                            if yeni_durum != row['durum']:
                                conn = db_baglan()
                                cur = conn.cursor()
                                cur.execute("UPDATE teklifler SET durum = ? WHERE id = ?", (yeni_durum, row['id']))
                                conn.commit()
                                conn.close()
                                st.success("Güncellendi!")
                                time.sleep(0.5)
                                st.rerun()

                            if c_git.button("Git ➡️", key=f"btn_git_{row['id']}"):
                                st.session_state.update(aktif_teklif_data=row, islem_turu="duzenle", sayfa_secimi="📝 Teklif Hazırla")
                                st.rerun()
            
            with tab_grafik:
                try:
                    onayli_grafik = en_guncel[en_guncel['durum'] == 'Onaylandı'].copy()
                    if not onayli_grafik.empty:
                        onayli_grafik['Yil'] = pd.to_datetime(onayli_grafik['tarih']).dt.year.astype(str)
                        chart_data = onayli_grafik.groupby(['Yil', 'para_birimi'])['toplam_tutar'].sum().unstack().fillna(0)
                        st.bar_chart(chart_data)
                    else:
                        st.info("Onaylanan proje bulunamadığı için grafik gösterilemiyor.")
                except:
                    st.error("Grafik oluşturulamadı.")
        else:
            st.info("Bu müşteriye ait henüz bir teklif kaydı bulunmuyor.")

    st.divider()

    # 3. VERİTABANI YÖNETİMİ (ALTA ALINDI ✅)
    st.subheader("🗄️ Veritabanı ve Müşteri Listesi")
    c_ara, c_yeni = st.columns([4, 1])
    with c_ara:
        arama_terimi = st.text_input("🔍 Müşteri Ara", placeholder="Tabloyu filtrelemek için firma adı yazın...", label_visibility="collapsed")
    with c_yeni:
        if st.button("➕ Yeni Müşteri", type="primary", use_container_width=True, key="new_cust_main"):
            ekleme_penceresi()

    if not df_musteriler.empty:
        if arama_terimi:
            df_goster = df_musteriler[df_musteriler['firma_adi'].str.contains(arama_terimi, case=False)]
        else:
            df_goster = df_musteriler
        st.dataframe(df_goster, hide_index=True, use_container_width=True)

elif st.session_state.sayfa_secimi == "⚙️ Sistem":
    st.title("⚙️ Sistem Yönetimi")
    st.markdown("Veritabanı yedekleme, geri yükleme ve genel uygulama ayarları.")
    
    tab_yedek, tab_ayarlar = st.tabs(["💾 Yedekleme & Kurtarma", "🛠️ Genel Ayarlar"])
    
    with tab_yedek:
        col_backup, col_restore = st.columns(2)
        
        # --- SOL TARAF: YEDEK ALMA ---
        with col_backup:
            st.info("### 📤 Yedek Al")
            st.write("Veritabanının kopyasını bilgisayarınızda güvenli bir klasöre kaydedin.")
            
            if st.button("Klasör Seç ve Yedekle", type="primary", use_container_width=True):
                # Tkinter penceresini gizle
                import tkinter as tk
                from tkinter import filedialog
                
                root = tk.Tk()
                root.withdraw()
                root.attributes('-topmost', True)
                
                klasor_yolu = filedialog.askdirectory(title="Yedeğin Kaydedileceği Klasörü Seçin")
                root.destroy()
                
                if klasor_yolu:
                    kaynak_dosya = "teklif_yonetim_sistemi.db"
                    if os.path.exists(kaynak_dosya):
                        zaman = datetime.now().strftime("%Y-%m-%d_%H-%M")
                        hedef = os.path.join(klasor_yolu, f"YEDEK_Saraks_{zaman}.db")
                        
                        try:
                            shutil.copy2(kaynak_dosya, hedef)
                            # Logo varsa onu da yedekle
                            if os.path.exists("logo.png"):
                                shutil.copy2("logo.png", os.path.join(klasor_yolu, "logo_yedek.png"))
                                
                            st.success(f"✅ Yedek Başarılı!\nKonum: {hedef}")
                        except Exception as e:
                            st.error(f"Hata: {e}")
                    else:
                        st.error("Veritabanı dosyası bulunamadı.")
        
        # --- SAĞ TARAF: GERİ YÜKLEME ---
        with col_restore:
            st.warning("### 📥 Geri Yükle")
            st.write("Daha önce aldığınız `.db` uzantılı yedek dosyasını yükleyerek sistemi o tarihe döndürün.")
            
            uploaded_db = st.file_uploader("Yedek Dosyası Seç (.db)", type="db")
            
            if uploaded_db:
                st.error("⚠️ DİKKAT: Bu işlem mevcut verilerin üzerine yazacaktır!")
                if st.button("🔴 ONAYLA VE YÜKLE", use_container_width=True):
                    try:
                        # Otomatik güvenlik yedeği al (ne olur ne olmaz)
                        if os.path.exists("teklif_yonetim_sistemi.db"):
                            shutil.copy2("teklif_yonetim_sistemi.db", "teklif_yonetim_sistemi.db.bak")
                        
                        # Dosyayı yaz
                        with open("teklif_yonetim_sistemi.db", "wb") as f:
                            f.write(uploaded_db.getbuffer())
                            
                        st.success("Sistem geri yüklendi! Lütfen sayfayı yenileyin.")
                        time.sleep(2)
                        st.rerun()
                    except Exception as e:
                        st.error(f"Hata: {e}")

    with tab_ayarlar:
        st.write("### 🔜 Gelecek Özellikler")
        st.info("Bu alana ileride şu özellikler eklenebilir:")
        st.markdown("""
        * Varsayılan KDV Oranı Ayarı
        * Varsayılan Para Birimi Ayarı
        * Logo Değiştirme / Yükleme Paneli
        * Renk Teması (Koyu / Açık Mod)
        """)

# ==============================================================================
# SAYFA 7: RESMİ SÖZLEŞME (TAM WORD + PDF KAYITLI)
# ==============================================================================
elif st.session_state.sayfa_secimi == "📜 Sözleşmeler":
    st.markdown("<h1 style='color:#8E44AD;'>📜 Resmi Sözleşme Masası</h1>", unsafe_allow_html=True)
    st.info("Sözleşmeler hem PDF hem de WORD (.docx) olarak 'Tam Metin' kaydedilir.")

    df_mus = musterileri_getir()
    
    if not df_mus.empty:
        c1, c2 = st.columns([2, 1])
        with c1:
            secilen_cari = st.selectbox("Sözleşme Yapılacak Müşteri", df_mus['firma_adi'].tolist(), index=None)
        
        if secilen_cari:
            cari_row = df_mus[df_mus['firma_adi'] == secilen_cari].iloc[0]
            conn = db_baglan()
            query = "SELECT * FROM teklifler WHERE musteri_id = ? AND durum = 'Onaylandı' ORDER BY id DESC"
            df_tek = pd.read_sql(query, conn, params=(int(cari_row['id']),))
            conn.close()

            if not df_tek.empty:
                df_tek['etiket'] = df_tek.apply(lambda x: f"{x['proje_adi']} | {format_para(x['toplam_tutar'], x['para_birimi'])}", axis=1)
                with c1:
                    secilen_proje_etiket = st.selectbox("Yayındaki Projeler", df_tek['etiket'].tolist(), index=None)
                
                if secilen_proje_etiket:
                    secilen_teklif = df_tek[df_tek['etiket'] == secilen_proje_etiket].iloc[0]
                    st.markdown("---")
                    
                    with st.form("sozlesme_form"):
                        st.markdown("### 👤 Müşteri Bilgileri")
                        cf1, cf2 = st.columns(2)
                        form_firma = cf1.text_input("Firma Unvanı", value=cari_row['firma_adi'])
                        form_sahis = cf2.text_input("Yetkili / Şahıs", value=cari_row.get('yetkili_kisi', ''))
                        
                        cf3, cf4 = st.columns([1, 2])
                        form_mus_vd = cf3.text_input("Vergi Dairesi / T.C. Kimlik no", value=str(cari_row.get('vergi_dairesi', '')))
                        form_mus_adres = cf4.text_input("Müşteri Adresi", value=cari_row['adres'])

                        st.markdown("### 📝 Şartlar")
                        ct1, ct2, ct3 = st.columns(3)
                        s_tarih = ct1.date_input("Sözleşme Tarihi", date.today())
                        b_tarih = ct2.date_input("İşin Bitiş Tarihi", date.today() + pd.Timedelta(days=30))
                        sehir = ct3.text_input("Sözleşme Yeri", value="Bursa")
                        
                        odeme_plani = st.text_area("Ödeme Şekli", placeholder="Toplam bedelin %50'si sipariş tarihinde, kalanı teslimattan önce tahsil edilir.")
                        
                        if st.form_submit_button("💾 KAYDET VE ARŞİVLE"):
                            try:
                                # İsim Mantığı
                                if form_firma.strip():
                                    final_mus_adi = f"{form_firma} (Yetkili: {form_sahis})" if form_sahis.strip() else form_firma
                                    klasor_ismi = form_firma.strip()
                                else:
                                    final_mus_adi = form_sahis
                                    klasor_ismi = form_sahis
                                
                                # Tarih Formatları
                                s_tarih_str = pd.to_datetime(str(s_tarih)).strftime('%d.%m.%Y')
                                b_tarih_str = pd.to_datetime(str(b_tarih)).strftime('%d.%m.%Y')

                                # 1. PDF ÜRET (Tam Metin)
                                pdf_data = create_contract_pdf(
                                    secilen_teklif['proje_adi'], secilen_teklif['toplam_tutar'], secilen_teklif['para_birimi'],
                                    final_mus_adi, form_mus_adres, form_mus_vd, s_tarih_str, b_tarih_str, sehir,
                                    "Binde Bir", 10, "2 Yıl", odeme_plani
                                )

                                # 2. WORD ÜRET (Tam Metin)
                                doc_object = create_contract_docx(
                                    secilen_teklif['proje_adi'], secilen_teklif['toplam_tutar'], secilen_teklif['para_birimi'],
                                    final_mus_adi, form_mus_adres, form_mus_vd, s_tarih_str, b_tarih_str, sehir,
                                    "Binde Bir", 10, "2 Yıl", odeme_plani
                                )

                                # Klasörleme
                                ana_dizin = "Sözleşmeler"
                                temiz_isim = "".join([x for x in klasor_ismi if x.isalnum() or x in ' -_']).strip()
                                mus_dizin = os.path.join(ana_dizin, temiz_isim)
                                if not os.path.exists(mus_dizin): os.makedirs(mus_dizin)
                                
                                # Dosya İsimleri
                                base_name = f"{s_tarih_str}_{secilen_teklif['proje_adi']}".replace(" ", "_")
                                pdf_yolu = os.path.join(mus_dizin, base_name + ".pdf")
                                docx_yolu = os.path.join(mus_dizin, base_name + ".docx")
                                
                                # KAYDET
                                with open(pdf_yolu, "wb") as f: f.write(pdf_data)
                                doc_object.save(docx_yolu)
                                
                                # Veritabanına PDF'i kaydet
                                sozlesme_kaydet(final_mus_adi, secilen_teklif['proje_adi'], s_tarih_str, pdf_yolu, format_para(secilen_teklif['toplam_tutar'], secilen_teklif['para_birimi']))
                                
                                st.success(f"✅ Sözleşmeler Hazır! Kayıt Yeri: {mus_dizin}")
                                
                                # İndirme Butonları
                                c_d1, c_d2 = st.columns(2)
                                b64_pdf = base64.b64encode(pdf_data).decode()
                                c_d1.markdown(f'<a href="data:application/pdf;base64,{b64_pdf}" download="{base_name}.pdf" style="display:block;width:100%;padding:10px;text-align:center;background:#E74C3C;color:white;border-radius:5px;text-decoration:none;">🔴 PDF İNDİR</a>', unsafe_allow_html=True)
                                
                                with open(docx_yolu, "rb") as f:
                                    docx_bytes = f.read()
                                b64_docx = base64.b64encode(docx_bytes).decode()
                                c_d2.markdown(f'<a href="data:application/vnd.openxmlformats-officedocument.wordprocessingml.document;base64,{b64_docx}" download="{base_name}.docx" style="display:block;width:100%;padding:10px;text-align:center;background:#2980B9;color:white;border-radius:5px;text-decoration:none;">🔵 WORD İNDİR</a>', unsafe_allow_html=True)

                                st.rerun()
                            except Exception as e:
                                st.error(f"Hata: {e}")
            else:
                st.warning(f"⚠️ {secilen_cari} adına 'Onaylandı' durumunda bir teklif bulunamadı.")
    else:
        st.warning("Henüz sisteme kayıtlı müşteri yok.")

    # --- ARŞİV ---
    st.markdown("---")
    st.subheader("🗄️ Sözleşme Arşivi")
    df_soz = sozlesmeleri_getir()
    if not df_soz.empty:
        for idx, row in df_soz.iterrows():
            with st.expander(f"📄 {row['tarih']} - {row['firma_adi']}"):
                ce1, ce2 = st.columns([3, 1])
                with ce1:
                    st.info(f"📁 Dosya: {row['dosya_yolu']}")
                    # Word kontrol
                    word_check = row['dosya_yolu'].replace(".pdf", ".docx")
                    if os.path.exists(word_check):
                        st.caption("✅ Word (.docx) yedeği mevcut.")
                    
                    u_firma = st.text_input("Müşteri", value=row['firma_adi'], key=f"uf_{row['id']}")
                    u_tutar = st.text_input("Tutar", value=row['tutar'], key=f"ut_{row['id']}")
                with ce2:
                    if st.button("🗑️ Sil", key=f"db_{row['id']}", use_container_width=True):
                        sozlesme_sil(row['id'])
                        st.rerun()
    else:
        st.info("Arşiv boş.")

               
# ==============================================================================
# SAYFA: TESLİM TUTANAĞI (NOTLAR BOŞ + WORD SÜTUNU + DOSYA SAATİ)
# ==============================================================================
elif st.session_state.sayfa_secimi == "🚛 Teslim Tutanağı":
    st.markdown("<h1 style='color:#2ECC71;'>🚛 Teslim ve Kabul Tutanağı</h1>", unsafe_allow_html=True)

    try:
        df_mus = musterileri_getir()
    except NameError:
        st.error("⚠️ HATA: Veritabanı fonksiyonları bulunamadı.")
        df_mus = pd.DataFrame()

    if not df_mus.empty:
        col1, col2 = st.columns([2, 1])
        with col1:
            secilen_cari = st.selectbox("Müşteri Seçiniz", df_mus['firma_adi'].tolist(), index=None)
        
        if secilen_cari:
            cari_row = df_mus[df_mus['firma_adi'] == secilen_cari].iloc[0]
            conn = db_baglan()
            query = "SELECT * FROM teklifler WHERE musteri_id = ? AND durum = 'Onaylandı' ORDER BY id DESC"
            df_tek = pd.read_sql(query, conn, params=(int(cari_row['id']),))
            conn.close()

            if not df_tek.empty:
                with col1:
                    df_tek['etiket'] = df_tek['proje_adi'] + " (Teklif ID: " + df_tek['id'].astype(str) + ")"
                    secilen_etiket = st.selectbox("Teslim Edilen Proje", df_tek['etiket'].tolist(), index=None)
                
                if secilen_etiket:
                    secilen_row = df_tek[df_tek['etiket'] == secilen_etiket].iloc[0]
                    secilen_teklif_id = int(secilen_row['id'])
                    proje_adi = secilen_row['proje_adi']

                    st.markdown("---")
                    
                    with st.form("tutanak_form"):
                        st.write("### 📅 Tarih Bilgileri")
                        c_d1, c_d2 = st.columns(2)
                        f_sozlesme_tarih = c_d1.date_input("Sözleşme Tarihi", date.today() - pd.Timedelta(days=30))
                        f_teslim_tarih = c_d2.date_input("İş Teslim Tarihi (Bugün)", date.today())
                        
                        st.markdown("### 📦 Ek-1: Teslim Edilen Ürün Listesi")
                        
                        # Otomatik Ürün Çekme
                        try:
                            gelen_urunler = teklif_urunlerini_getir(secilen_teklif_id)
                        except:
                            gelen_urunler = pd.DataFrame()

                        if gelen_urunler.empty:
                            st.warning("Teklifte kayıtlı ürün yok. Lütfen elle doldurun.")
                            gelen_urunler = pd.DataFrame([{"Kod": "", "Urun": "", "Adet": ""}])
                        
                        # --- DEĞİŞİKLİK BURADA: ARTIK VARSAYILAN DEĞER BOŞ ---
                        if "Not" not in gelen_urunler.columns:
                            gelen_urunler["Not"] = "" # Eskiden "Eksiksiz Teslim Edildi" yazıyordu

                        # Tablo
                        edited_df = st.data_editor(
                            gelen_urunler[["Kod", "Urun", "Adet", "Not"]],
                            num_rows="dynamic",
                            use_container_width=True,
                            column_config={
                                "Kod": st.column_config.TextColumn("Kodu", width="small"),
                                "Urun": st.column_config.TextColumn("Ürün Adı", width="large"),
                                "Adet": st.column_config.TextColumn("Adet", width="small"),
                                "Not": st.column_config.TextColumn("Teslim Notu", width="medium")
                            }
                        )
                        
                        st.markdown("---")
                        
                        if st.form_submit_button("✅ TUTANAĞI OLUŞTUR"):
                            try:
                                soz_tarih_str = pd.to_datetime(str(f_sozlesme_tarih)).strftime('%d.%m.%Y')
                                tes_tarih_str = pd.to_datetime(str(f_teslim_tarih)).strftime('%d.%m.%Y')
                                
                                clean_df = edited_df[edited_df["Urun"].str.strip() != ""].copy()
                                
                                # Sütun İsmi Düzeltme (Word ve PDF hatasını önler)
                                clean_df = clean_df.rename(columns={"Urun": "Ürün Adı"})
                                
                                # PDF ve Word
                                pdf_data = create_delivery_pdf(secilen_cari, proje_adi, soz_tarih_str, tes_tarih_str, clean_df)
                                doc_object = create_delivery_docx(secilen_cari, proje_adi, soz_tarih_str, tes_tarih_str, clean_df)
                                
                                # Kayıt
                                ana_klasor = "Teslim_Tutanaklari"
                                temiz_isim = "".join([x for x in secilen_cari if x.isalnum() or x in ' -_']).strip()
                                mus_klasor = os.path.join(ana_klasor, temiz_isim)
                                if not os.path.exists(mus_klasor): os.makedirs(mus_klasor)
                                
                                # --- DOSYA İSMİNE SAAT EKLENDİ (HATA ÇÖZÜMÜ) ---
                                zaman_damgasi = datetime.now().strftime("%H%M%S")
                                dosya_ismi = f"Tutanak_{proje_adi}_{tes_tarih_str}_{zaman_damgasi}".replace(" ", "_")
                                
                                pdf_yolu = os.path.join(mus_klasor, dosya_ismi + ".pdf")
                                docx_yolu = os.path.join(mus_klasor, dosya_ismi + ".docx")
                                
                                with open(pdf_yolu, "wb") as f: f.write(pdf_data)
                                doc_object.save(docx_yolu)
                                
                                tutanak_kaydet(secilen_cari, proje_adi, tes_tarih_str, pdf_yolu)
                                
                                st.success(f"✅ Tutanak Oluşturuldu! Kayıt: {mus_klasor}")
                                
                                # İndirme Butonları
                                d1, d2 = st.columns(2)
                                b64_pdf = base64.b64encode(pdf_data).decode()
                                d1.markdown(f'<a href="data:application/pdf;base64,{b64_pdf}" download="{dosya_ismi}.pdf" style="display:block;width:100%;padding:10px;text-align:center;background:#E74C3C;color:white;border-radius:5px;text-decoration:none;">🔴 PDF İNDİR</a>', unsafe_allow_html=True)
                                
                                with open(docx_yolu, "rb") as f: docx_bytes = f.read()
                                b64_docx = base64.b64encode(docx_bytes).decode()
                                d2.markdown(f'<a href="data:application/vnd.openxmlformats-officedocument.wordprocessingml.document;base64,{b64_docx}" download="{dosya_ismi}.docx" style="display:block;width:100%;padding:10px;text-align:center;background:#2980B9;color:white;border-radius:5px;text-decoration:none;">🔵 WORD İNDİR</a>', unsafe_allow_html=True)

                            except Exception as e:
                                st.error(f"Oluşturma Hatası: {e}")
            else:
                st.warning("Bu müşterinin 'Onaylandı' durumunda bir projesi yok.")
    else:
        st.warning("Kayıtlı müşteri yok.")

    # --- ARŞİV ---
    st.markdown("---")
    st.subheader("🗄️ Tutanak Arşivi")
    try:
        df_tut = tutanaklari_getir()
        if not df_tut.empty:
            for idx, row in df_tut.iterrows():
                with st.expander(f"🚛 {row['tarih']} - {row['firma_adi']}"):
                    c1, c2 = st.columns([3, 1])
                    with c1: st.write(f"📁 `{row['dosya_yolu']}`")
                    with c2:
                        if st.button("🗑️ Sil", key=f"del_tut_{row['id']}", use_container_width=True):
                            tutanak_sil(row['id'])
                            st.rerun()
        else:
            st.info("Arşiv boş.")
    except NameError:
        st.error("Veritabanı fonksiyonu eksik.")
